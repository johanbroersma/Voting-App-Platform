#!/usr/bin/env python3
"""
Board Voting App — Election Officer Manual
Generates board_manual.docx   (run: python3 build_board_manual.py)
Hosted solution edition — for use with votingapp.ca managed hosting.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

NAVY   = RGBColor(0x1a, 0x27, 0x44)
GOLD   = RGBColor(0xb8, 0x94, 0x2a)
GREEN  = RGBColor(0x2d, 0x6a, 0x4f)
RED    = RGBColor(0x9b, 0x23, 0x35)
GREY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xff, 0xff, 0xff)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

def _space(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def shade_cell(cell, hex_colour):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph(); _space(p, 18, 6)
    p.paragraph_format.page_break_before = True
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '1a2744')
    pBdr.append(bot); pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph(); _space(p, 12, 4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

def h3(text):
    p = doc.add_paragraph(); _space(p, 8, 3)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREY

def body(text):
    p = doc.add_paragraph(); _space(p, 0, 5)
    r = p.add_run(text); r.font.size = Pt(11)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet'); _space(p, 0, 3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.add_run(text).font.size = Pt(11)

def numbered(text):
    p = doc.add_paragraph(style='List Number'); _space(p, 0, 3)
    p.paragraph_format.left_indent = Cm(0.5)
    p.add_run(text).font.size = Pt(11)

def callout(label, text, bg_hex, border_hex):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0); shade_cell(cell, bg_hex)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '12' if side in ('top', 'left') else '4')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), border_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)
    lp = cell.paragraphs[0]; _space(lp, 2, 2)
    lr = lp.add_run(f'{label}  '); lr.bold = True; lr.font.size = Pt(10)
    lp.add_run(text).font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def note(text):    callout('NOTE:',            text, 'E8F0FA', '1A3A5C')
def tip(text):     callout('TIP:',             text, 'F0FAF0', '2D6A4F')
def warning(text): callout('IMPORTANT:',       text, 'FFF4E0', 'B8942A')
def danger(text):  callout('ACTION REQUIRED:', text, 'FDF0F0', '9B2335')

def table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]; shade_cell(cell, '1A2744')
        p = cell.paragraphs[0]; _space(p, 2, 2)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE
    for ri, row_data in enumerate(rows):
        bg = 'F8F6F1' if ri % 2 else 'FFFFFF'
        for ci, text in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]; shade_cell(cell, bg)
            p = cell.paragraphs[0]; _space(p, 2, 2)
            p.add_run(str(text)).font.size = Pt(10)
    if col_widths:
        for row in tbl.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
cover = doc.add_paragraph(); cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
_space(cover, 80, 6)
r = cover.add_run('Board Voting App'); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; _space(sub, 6, 10)
r = sub.add_run('Election Officer Manual'); r.font.size = Pt(16); r.font.color.rgb = GOLD

ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER; _space(ver, 0, 30)
r = ver.add_run('Hosted Solution Edition — votingapp.ca'); r.font.size = Pt(11); r.font.color.rgb = GREY

for line in [
    'Board Elections · Motion Voting',
    'For Boards, Committees, and Organizations',
    '',
    'Confidential — For Election Officers & Chairman Use Only',
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _space(p, 0, 4)
    r = p.add_run(line); r.font.size = Pt(11)
    r.bold = 'Confidential' in line
    r.font.color.rgb = NAVY if 'Confidential' in line else GREY

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
toc = doc.add_paragraph(); _space(toc, 0, 8)
r = toc.add_run('Table of Contents'); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

toc_entries = [
    ('1.',     'System Overview'),
    ('2.',     'Accessing the App'),
    ('3.',     'App Settings'),
    ('   3a.', 'Organization Identity'),
    ('   3b.', 'Passwords'),
    ('4.',     'Home Screen'),
    ('5.',     'Election Setup'),
    ('   5a.', 'Election Details'),
    ('   5b.', 'Nominees'),
    ('   5c.', 'Settings — Completion Message, Voter URL & Danger Zone'),
    ('6.',     'Voter Tokens'),
    ('7.',     'Running the Election — Round Control'),
    ('   7a.', 'Starting a Round'),
    ('   7b.', 'Opening & Closing Voting'),
    ('   7c.', 'Monitoring Participation'),
    ('   7d.', 'Round Transition'),
    ('   7e.', 'Completing the Election'),
    ('8.',     'Paper Ballot Entry'),
    ('9.',     'The Voter Page'),
    ('10.',    'Election Dashboard'),
    ('11.',    'Motion Vote'),
    ('   11a.','Configuring the Motion'),
    ('   11b.','Motion Controls'),
    ('   11c.','Motion Results'),
    ('   11d.','Paper Ballot Entry — Motion Vote'),
    ('12.',    'Roles & Responsibilities'),
    ('13.',    'Meeting Day Checklist'),
    ('14.',    'Suggested Order of Events'),
    ('15.',    'Troubleshooting'),
    ('16.',    'Quick Reference'),
]
for num, title in toc_entries:
    p = doc.add_paragraph(); _space(p, 0, 3)
    indent = len(num) - len(num.lstrip())
    p.paragraph_format.left_indent = Cm(indent * 0.2)
    top = not num.startswith('   ')
    r1 = p.add_run(num + '  '); r1.font.size = Pt(11); r1.bold = top
    r1.font.color.rgb = NAVY if top else GREY
    r2 = p.add_run(title); r2.font.size = Pt(11)
    r2.font.color.rgb = NAVY if top else GREY

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ══════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════
h1('1.  System Overview')

body(
    'The Board Voting App is a browser-based voting platform for boards, committees, and '
    'member organizations. Your instance is hosted and managed on votingapp.ca — there is '
    'no software to install and no server to run. You access the admin hub through your web '
    'browser, and voters use the ballot page on their own phones or tablets.'
)

h2('What the App Supports')
for item in [
    'Multi-round board elections — configure nominees, open positions, and votes per voter',
    'Motion Vote — a one-round vote on a motion or resolution with a configurable majority threshold',
    'Digital ballots via personal phones using unique 4-digit token codes',
    'Paper ballot entry for members who prefer to vote on paper',
    'Absentee vote support for both the election and the motion vote',
    'Real-time vote count updates — Round Control and dashboards refresh every 3 seconds',
    'One token card per voter covers all election rounds and the motion vote',
    'Password-protected access levels so each role sees only what they need',
]:
    bullet(item)

h2('The Three Pages')
table(
    ['Page', 'Who Uses It', 'Purpose'],
    [
        ['Admin Hub\n(your app URL/)', 'Election officer, chairman, volunteers', 'All setup, round control, motion vote, results, token printing'],
        ['Voter Page\n(your app URL/vote.html)', 'All eligible voters', 'Enter token, submit election ballot, cast motion vote'],
        ['App Settings\n(accessible from Admin Hub)', 'Election officer only', 'Organization identity and all passwords'],
    ],
    col_widths=[4.5, 4.5, 7.0]
)

note('Your specific app URL (e.g. https://myboard.votingapp.ca) is provided by your hosting administrator. All pages are accessed via that same base URL.')


# ══════════════════════════════════════════════════════════════════
# 2. ACCESSING THE APP
# ══════════════════════════════════════════════════════════════════
h1('2.  Accessing the App')

body(
    'The Board Voting App runs entirely in a web browser. No installation is required on any device.'
)

h2('Admin Hub')
numbered('Open your web browser (Chrome, Safari, Edge, or Firefox).')
numbered('Navigate to your app URL — e.g. https://myboard.votingapp.ca')
numbered('Enter the app password on the login screen and tap Log In.')
numbered('The Home Screen appears, showing the navigation tiles.')

h2('Voter Devices')
numbered('Voters open the voter page in any browser on their phone or tablet.')
numbered('They can type the voter URL directly, or scan the QR code printed on their token card.')
numbered('No app installation, no account, no login — just the 4-digit token code.')

note('The voter URL is your app URL with /vote.html at the end (e.g. https://myboard.votingapp.ca/vote.html). It is shown in the Voter Tokens screen and printed on every token card.')

warning(
    'The app requires an active internet connection on all devices. Voters must be able to reach '
    'your app URL from their phones. Confirm this works before the meeting by testing on a phone.'
)

h2('Device Requirements')
table(
    ['Device', 'Requirements'],
    [
        ['Admin device (election officer)', 'Any modern browser. Keep this tab open throughout the meeting.'],
        ['Chairman device', 'Any browser — used for Election Dashboard only.'],
        ['Paper ballot station', 'Any browser — used for Paper Ballot Entry only. Dedicate a separate device.'],
        ['Voter phones/tablets', 'Any browser. Mobile data or Wi-Fi — either works if it can reach the app URL.'],
    ],
    col_widths=[5.0, 11.0]
)


# ══════════════════════════════════════════════════════════════════
# 3. APP SETTINGS
# ══════════════════════════════════════════════════════════════════
h1('3.  App Settings')

body(
    'App Settings is where you configure your organization\'s identity and manage all passwords. '
    'Access it from the Home Screen by clicking the App Settings tile. '
    'The admin password is required.'
)

h2('3a.  Organization Identity')
body('Fill in the following fields so the app reflects your organization:')
table(
    ['Field', 'Description', 'Example'],
    [
        ['Organization Name', 'Displayed on the voter page header and token cards.', 'NRCEA School Board'],
        ['Election Name', 'Subtitle shown on the home screen and voter page.', 'Board of Directors Election 2026'],
        ['Organization Logo', 'Upload your logo (optional). Appears on the home screen. Square image recommended.', '(upload image)'],
    ],
    col_widths=[4.0, 8.5, 3.5]
)
body('Click Save Identity Settings to apply all identity changes.')
tip('Set your organization name and election name before printing token cards — the name appears on every card.')

h2('3b.  Passwords')
body(
    'The app uses six independent passwords so each role only accesses what they need. '
    'Change all passwords from their defaults before the meeting.'
)
table(
    ['Password', 'Default', 'Protects', 'Held By'],
    [
        ['App Password',         'votevote2024',   'Admin hub login screen — entry gate for all officers', 'All officers'],
        ['Admin Password',       'churchvoting',   'App Settings',                                          'Election Officer only'],
        ['Election Password',    'election2024',   'Election Setup, Round Control, and Motion Vote Control', 'Election Officer only'],
        ['Results Password',     'results2024',    'Election Dashboard (read-only results)',                'Chairman only'],
        ['Paper Ballot Password','paperentry2024', 'Paper Ballot Entry station',                            'Paper ballot volunteer'],
        ['Tokens Password',      'tokens2024',     'Voter Tokens section',                                  'Token administrator'],
    ],
    col_widths=[3.8, 3.0, 5.2, 4.0]
)
danger(
    'Change ALL passwords before the meeting. The defaults above are publicly documented. '
    'Go to App Settings to change each one. Record your chosen passwords securely and '
    'distribute each password only to the person who needs it.'
)
body(
    'To change a password: enter the current password in the "Current Password" field, '
    'type the new password twice, and click the Update button for that password. '
    'If you are still using the default, leave the current password field blank.'
)
note(
    'All password fields have autocapitalize, autocorrect, and spellcheck disabled to prevent '
    'mobile keyboard interference. Type carefully — do not copy-paste from other apps.'
)


# ══════════════════════════════════════════════════════════════════
# 4. HOME SCREEN
# ══════════════════════════════════════════════════════════════════
h1('4.  Home Screen')

body(
    'After logging in with the app password, the Home Screen shows the navigation tiles. '
    'A yellow warning banner appears at the top if the default app password is still in use.'
)

table(
    ['Tile', 'Destination', 'Password Required'],
    [
        ['Board Election',      'Election Hub — Setup, Round Control, Dashboard',    'Election password (at hub)'],
        ['Paper Ballot Entry',  'Paper ballot station for election and motion votes', 'Paper ballot password'],
        ['Voter Tokens',        'Token generation, printing, and QR code',           'Tokens password'],
        ['Motion Vote',         'Motion Vote Control — configure and run a vote',    'Election password'],
        ['App Settings',        'Organization identity and all passwords',           'Admin password'],
    ],
    col_widths=[4.0, 7.0, 5.0]
)

tip('The election officer tab should stay on the Election Hub or Motion Vote Control throughout the meeting. Use a separate browser tab for the Election Dashboard (chairman) and Paper Ballot Entry (volunteer).')


# ══════════════════════════════════════════════════════════════════
# 5. ELECTION SETUP
# ══════════════════════════════════════════════════════════════════
h1('5.  Election Setup')

body(
    'Election Setup is used before the meeting to configure the board election. '
    'From the Home Screen, click Board Election. In the Election Hub, click Election Setup. '
    'The election password is required.'
)
body('Election Setup has three tabs: Details, Nominees, and Settings.')

h2('5a.  Election Details')
body('Click the Details tab and fill in:')
table(
    ['Field', 'Description', 'Example'],
    [
        ['Expected Voters',          'Eligible voters attending in person. Used for majority calculation and participation %.', '35'],
        ['Absentee Votes (Round 1)', 'Pre-meeting paper ballots received. Added to the majority base in Round 1 only.', '3'],
    ],
    col_widths=[4.5, 7.5, 4.0]
)
body(
    'Below these fields the system shows the Majority Required threshold. '
    'Majority = floor(n/2) + 1. Round 1 uses Expected Voters + Absentee Votes. '
    'Round 2 and later use Expected Voters only.'
)
warning('Enter the absentee count before the meeting. Each absentee ballot is entered individually during Round 1 via Paper Ballot Entry, with the Absentee checkbox ticked.')

h2('5b.  Nominees')
body('Click the Nominees tab and fill in:')
table(
    ['Field', 'Description', 'Example'],
    [
        ['Nominees',             'One name per line. All candidates nominated for the election.', 'Jane Smith\nMark Johnson\nPriya Patel'],
        ['Open Positions',       'Number of seats to fill in this election.',                     '2'],
        ['Votes per Voter (R1)', 'How many candidates each voter may select in Round 1. Typically equals the number of positions.', '2'],
    ],
    col_widths=[4.5, 8.5, 3.0]
)
body('Click Save Nominees & Setup when done. A green confirmation appears.')
warning('Saving nominees resets all in-progress election data (ballots, results). Never click Save during an active voting round.')

h3('Printing Paper Ballots')
bullet('Enter the number of paper ballots needed and click Print Ballots.')
bullet('Each ballot shows the organization name, positions to fill, and nominee names with checkboxes.')
note('The same printed ballots can be reused in all rounds. The chairman verbally announces which candidates remain on the ballot each round.')

h2('5c.  Settings — Completion Message, Voter URL & Danger Zone')
body('Click the Settings tab:')

h3('Election Completion Message')
body(
    'Enter a custom message shown to voters and on the Election Dashboard when the election is complete. '
    'Optionally add a Feedback Survey Link — if provided, a "Complete Feedback Survey" button '
    'appears on the completion screen so voters can access a post-election survey.'
)

h3('Voter Page URL')
body(
    'If your app URL is not automatically detected, enter the voter page URL here (your app URL + /vote.html). '
    'This URL is shown on token cards and the Voter Tokens screen. '
    'Click Save URL if you make any changes.'
)
warning('Confirm the voter URL before printing token cards. The URL and QR code printed on every card come from this saved value.')

h3('Reset All Election Data')
body(
    'The red Reset All Election Data button (Danger Zone) permanently erases all nominees, '
    'ballots, tokens, results, and resets passwords to defaults. Two confirmation prompts prevent '
    'accidental activation.'
)
danger('Reset All Election Data cannot be undone. Use only to prepare for a fresh election — never during a meeting.')


# ══════════════════════════════════════════════════════════════════
# 6. VOTER TOKENS
# ══════════════════════════════════════════════════════════════════
h1('6.  Voter Tokens')

body(
    'Each eligible voter receives one token card before the meeting. The 4-digit code on '
    'the card is their key to vote in all election rounds and the motion vote — '
    'they never need a new card.'
)
body('Access from the Home Screen: click Voter Tokens and enter the tokens password.')

h2('Generating Tokens')
numbered('Enter the total number of eligible voters (include both in-person and absentee).')
numbered('Click Generate Tokens. The system creates unique 4-digit codes.')
numbered('The token grid displays all codes. Used tokens are shown with a grey strikethrough.')
warning('Generate tokens only once before the meeting. Regenerating during a live election invalidates all distributed cards.')

h2('Printing Token Cards')
body(
    'Click Print Token Cards to open the print view. Each card shows: organization name, '
    'the 4-digit code in large text, a QR code linking to the voter page, and the voter URL. '
    'Print on card stock and cut into individual cards. Distribute one per eligible voter at the entrance.'
)
tip('Print and cut cards the day before the meeting. Use A4 card stock for best results. Have a door steward distribute cards only to verified eligible voters.')

h2('Voter URL & QR Code Display')
body(
    'The Voter Tokens screen also displays the current voter URL and a large QR code. '
    'Scan the QR code to confirm it opens the correct voter page before the meeting.'
)
note('If the voter URL is wrong on this screen, update it in Election Setup → Settings → Voter Page URL, then return here to verify.')

h2('Changing the Tokens Password')
body('Enter the current tokens password, type the new password twice, and click Update Tokens Password.')


# ══════════════════════════════════════════════════════════════════
# 7. RUNNING THE ELECTION — ROUND CONTROL
# ══════════════════════════════════════════════════════════════════
h1('7.  Running the Election — Round Control')

body(
    'Round Control is used during the meeting to manage the board election. '
    'From the Election Hub, click Round Control. The election password is required.'
)
body(
    'Four stat boxes update live every 3 seconds. The control bar has three buttons. '
    'A "▶ Round Control" shortcut button appears in the Election Hub header once nominees are saved.'
)

h2('7a.  Starting a Round')
body(
    'When no election is in progress, Round Control shows a "Ready to Start" panel '
    'listing nominees and expected voters. '
    'Click Start Round 1 Voting to begin.'
)
tip('Do not click Start Round 1 until the chairman has introduced the candidates. Once clicked, the ballot is immediately available on voters\' phones.')

h2('7b.  Opening & Closing Voting')
table(
    ['Button', 'Action', 'When to Use'],
    [
        ['▶ Open Voting',  'Opens voting — members can now submit ballots.', 'At the start of each voting period, after the chairman\'s announcement.'],
        ['■ Close Voting', 'Closes voting — no further ballots accepted.', 'When participation is sufficient, before processing results.'],
        ['📊 End Round',   'Confirms the round and opens the Round Transition screen.', 'After voting is closed. Requires voting to be closed first.'],
    ],
    col_widths=[3.5, 6.0, 6.5]
)
warning('End Round requires voting to be closed first. A confirmation dialog shows the ballot count — click Confirm to proceed.')
body('When voting is closed, all voter screens update within 3 seconds to show "Voting Round Closed — please wait for further instructions."')

h2('7c.  Monitoring Participation')
body('The stat boxes update every 3 seconds:')
table(
    ['Statistic', 'Meaning'],
    [
        ['Ballots In',      'Total ballots this round (digital + paper combined).'],
        ['Paper Ballots',   'Ballots entered via the Paper Ballot Entry station.'],
        ['Absentee',        'Paper ballots marked as absentee (Round 1 only).'],
        ['Expected Voters', 'Round 1: attending + absentee. Round 2+: attending only.'],
        ['Participation %', 'Ballots In ÷ Expected Voters × 100. Shown green when ≥ 50%.'],
    ],
    col_widths=[4.5, 11.5]
)
note('Live vote counts per candidate are visible only on Round Control and the Election Dashboard — not on any public screen or voter phone.')

h2('7d.  Round Transition')
body('After clicking End Round and confirming, the Round Transition screen opens with four sections:')
for step in [
    'Final Results — all candidates ranked by votes received.',
    'Step 1 — Mark as Elected: candidates who reached majority are pre-checked. Adjust checkboxes if needed.',
    'Step 2 — Select Who Advances: choose which remaining candidates go to the next round (all advance by default). Uncheck to exclude.',
    'Step 3 — Next Round Settings: set votes per voter for the next round. Auto-filled with remaining positions.',
    'If all positions are filled: a green banner appears. Steps 2 and 3 are hidden. Only Complete This Election is available.',
    'Click Launch Next Round (opens Round Control with voting closed so the chairman can announce first) OR Complete This Election.',
]:
    numbered(step)

h2('7e.  Completing the Election')
body('When all positions are filled, click Complete This Election. The system:')
bullet('Records all elected candidates.')
bullet('Shows the Election Complete screen in Round Control.')
bullet('All voter phones display the election completion message simultaneously.')
body(
    'At this point, if a Motion Vote is configured, the election officer can switch to '
    'Motion Vote Control from the Home Screen to run the motion vote.'
)


# ══════════════════════════════════════════════════════════════════
# 8. PAPER BALLOT ENTRY
# ══════════════════════════════════════════════════════════════════
h1('8.  Paper Ballot Entry')

body(
    'The Paper Ballot Entry screen is used by a dedicated volunteer to enter votes for '
    'members who voted on paper. It is accessed from the Home Screen — the paper ballot '
    'password is required. The screen automatically detects whether an election round or '
    'motion vote is open and shows the correct form.'
)
warning('Open Paper Ballot Entry on a dedicated device or a separate browser tab. It is operated by a volunteer, not the election officer.')

h2('Entering an Election Ballot')
numbered('Confirm the correct round is shown at the top of the screen.')
numbered('Click each candidate\'s name from the paper ballot. Selected candidates are highlighted with a checkmark.')
numbered('(Round 1 only) If the ballot is from an absentee voter, tick the "Absentee ballot" checkbox before submitting.')
numbered('Click Submit Paper Ballot. The vote is recorded and the form resets.')
numbered('The submitted ballot appears in the log at the bottom of the screen.')

h2('Editing or Deleting a Ballot')
bullet('To edit: click the ✏ Edit button next to a ballot in the log. Make corrections and click Update Paper Ballot.')
bullet('To delete the most recent entry: click Delete Last Entry (red button) and confirm.')
note('Only the most recently submitted ballot can be deleted. Use the Edit button to correct earlier entries.')

h2('When No Voting Is Open')
body(
    'If no election round or motion vote is currently open, the screen shows '
    '"No election or motion vote is currently open." The form appears automatically within '
    '4 seconds when voting opens.'
)


# ══════════════════════════════════════════════════════════════════
# 9. THE VOTER PAGE
# ══════════════════════════════════════════════════════════════════
h1('9.  The Voter Page')

body(
    'Voters use a single URL throughout the entire meeting — for both the board election '
    'and the motion vote. The page automatically detects what is active and shows the '
    'correct ballot. No navigation or page refreshing is needed.'
)

h2('Voting Flow — Board Election')
numbered('Voter opens the voter URL or scans the QR code on their token card.')
numbered('Voter enters their 4-digit token code and taps Submit Token.')
numbered('The ballot appears: round number, candidates, and how many to select.')
numbered('Voter taps each candidate to select. A coloured highlight confirms the choice.')
numbered('Voter taps Submit My Vote. A "Vote Submitted!" confirmation appears.')
numbered('If another round opens, a notification appears — voter taps Vote Now.')
note('The same token card is used for all rounds. Voters never need to re-enter their token for a new round of the same election.')

h2('Voting Flow — Motion Vote')
numbered('If the voter already voted in the election, the page transitions automatically — no re-entry of the token is needed.')
numbered('If no election was held, the voter enters their token as usual.')
numbered('The motion ballot appears with the motion text and answer options.')
numbered('Voter selects their answer and taps Cast My Vote.')
numbered('A "Vote Recorded" confirmation shows the motion text and selected answer.')

h2('All Voter Screen States')
table(
    ['State', 'When It Appears'],
    [
        ['Token Entry',       'Voting is open; voter has not yet entered their token this session.'],
        ['Election Ballot',   'Token entered; election round is open.'],
        ['Motion Ballot',     'Token entered; motion vote is open.'],
        ['Vote Submitted',    'Election ballot confirmed. Page polls for next round.'],
        ['Vote Recorded',     'Motion vote confirmed.'],
        ['Voting Closed',     'Round or motion closed while voter is on ballot screen.'],
        ['Waiting',           'Voting not yet open or between rounds.'],
        ['Election Complete', 'All positions filled — election completion message shown.'],
        ['Vote Complete',     'Motion vote marked complete.'],
    ],
    col_widths=[5.0, 11.0]
)
warning('Votes are final once submitted. Digital ballots cannot be changed. Only paper ballots can be edited via the Paper Ballot Entry station.')


# ══════════════════════════════════════════════════════════════════
# 10. ELECTION DASHBOARD
# ══════════════════════════════════════════════════════════════════
h1('10.  Election Dashboard')

body(
    'The Election Dashboard is a read-only view for the chairman showing live vote counts '
    'and election status, refreshing every 3 seconds. '
    'From the Election Hub, click Election Dashboard and enter the results password.'
)

h2('What the Dashboard Shows')
table(
    ['Section', 'Content'],
    [
        ['Status Card',     'Current round, voting status pill (Open/Closed/Configuring), majority required, votes per voter, candidates on ballot.'],
        ['Info Strip',      'Expected Voters and Absentee Votes pills.'],
        ['Results Section', 'All candidates with proportional vote bars, vote counts, and ★ Majority / ✓ Elected badges.'],
    ],
    col_widths=[4.5, 11.5]
)

body('The results section adapts to the election state:')
bullet('Before Round 1: shows nominees, positions, and Round 1 majority threshold.')
bullet('During a round: shows live vote counts and bars labelled "● Voting Open".')
bullet('After End Round (before transition): shows a "Round N Results — Awaiting transition" card.')
bullet('After transition: shows committed historical results for all completed rounds.')
tip('The "Awaiting transition" card lets the chairman review results and prepare their announcement before the election officer processes the transition.')
warning('The dashboard shows actual vote counts. Keep the results password confidential and never leave the dashboard visible on an unattended screen.')


# ══════════════════════════════════════════════════════════════════
# 11. MOTION VOTE
# ══════════════════════════════════════════════════════════════════
h1('11.  Motion Vote')

body(
    'The Motion Vote is a one-round vote on a motion or resolution from the membership. '
    'It can be run before, after, or independently of the board election. '
    'From the Home Screen, click Motion Vote and enter the election password.'
)

h2('11a.  Configuring the Motion')
body('The Motion Vote Control screen handles both setup and control. When no motion is active, fill in:')
table(
    ['Field', 'Description', 'Example'],
    [
        ['Motion Text',        'The motion or question members will vote on.',             'Do you approve the proposed budget amendment?'],
        ['Answer Options',     'The choices members select. At least two required. Default: Yes / No. Add or remove as needed.', 'Yes / No'],
        ['Expected Voters',    'Eligible voters for this motion. Defaults to the election expected voters if left at 0.', '35'],
        ['Votes per Voter',    'How many options each voter can select. Usually 1.',       '1'],
        ['Majority Threshold', '50%+1 (simple majority), 75% (supermajority), or a custom number of votes required.', '50%+1'],
    ],
    col_widths=[4.0, 8.0, 4.0]
)
body('Click Open Motion Voting to save the configuration and open voting simultaneously.')
warning('Once motion voting is opened with a configuration, the setup fields are locked. To change the motion, you must close voting first and reset (if votes have been cast, resetting clears all votes).')

h2('11b.  Motion Controls')
table(
    ['Button', 'Action', 'Notes'],
    [
        ['▶ Open Voting',      'Opens motion voting — members can now cast their vote.', 'Voter phones switch to the motion ballot automatically.'],
        ['■ Close Voting',     'Closes motion voting.',                                   'Voter phones update within 3 seconds.'],
        ['✅ Mark Complete',   'Finalises the motion vote permanently.',                  'Disabled while voting is open. Requires confirmation.'],
        ['🔄 Reset & New Vote','Clears all motion votes and returns to configuration.',    'Use only after the current motion is complete or to start over.'],
    ],
    col_widths=[3.8, 5.7, 6.5]
)

body('The stat boxes update every 3 seconds:')
table(
    ['Statistic', 'Meaning'],
    [
        ['Votes Cast', 'Total votes received (digital + paper).'],
        ['Expected',   'Expected voters for this motion.'],
        ['Turnout %',  'Votes Cast ÷ Expected × 100. Shown green when ≥ 50%.'],
        ['Majority',   'Minimum votes needed to pass based on the configured threshold.'],
    ],
    col_widths=[4.0, 12.0]
)
warning('When the motion vote is marked complete, voter phones show: "Thank you for your vote. The voting is now complete." The election officer should announce results verbally from the live results panel on this screen.')

h2('11c.  Motion Results')
body(
    'Live results are shown on the Motion Vote Control screen while voting is open or after it closes. '
    'For each answer option, the screen shows: vote count, proportional bar, and percentage. '
    'A green ✓ Majority badge appears when the threshold is reached.'
)
note('The Motion Vote does not have a separate dashboard — only the election officer (on Motion Vote Control) can see live motion results. There is no chairman results view for the motion vote.')
tip('The election officer can describe the result verbally to the chairman, who can then announce it. Alternatively, the election officer can share their screen briefly for the announcement.')

h2('11d.  Paper Ballot Entry — Motion Vote')
body(
    'The Paper Ballot Entry screen automatically switches to motion vote mode when a motion vote is open.'
)
numbered('The screen shows the motion text and answer options.')
numbered('Click the member\'s selected answer to highlight it.')
numbered('Click Submit Paper Vote. The vote is recorded and the form resets.')
note('The motion vote paper ballot does not have an absentee checkbox — all motion votes are treated equally regardless of how they were cast.')


# ══════════════════════════════════════════════════════════════════
# 12. ROLES & RESPONSIBILITIES
# ══════════════════════════════════════════════════════════════════
h1('12.  Roles & Responsibilities')

table(
    ['Role', 'Responsibilities', 'Passwords Needed'],
    [
        ['Election Officer',       'Operates Election Setup, Round Control, and Motion Vote Control; opens/closes voting; processes round transitions; announces motion results.', 'App, Admin, Election'],
        ['Chairman',               'Announces candidates and election results; accesses Election Dashboard; does not operate any control functions.', 'App, Results'],
        ['Paper Ballot Volunteer', 'Operates Paper Ballot Entry; collects and enters paper ballots for both election and motion vote.', 'App, Paper Ballot'],
        ['Token Administrator',    'Generates and prints token cards before the meeting; distributes them at the door.', 'App, Tokens'],
        ['Door Steward',           'Distributes token cards to verified eligible voters only.', 'App (for login)'],
    ],
    col_widths=[3.8, 8.2, 4.0]
)

warning('Never share the results password with the election officer, or the admin/election password with the chairman. Keep roles separated.')


# ══════════════════════════════════════════════════════════════════
# 13. MEETING DAY CHECKLIST
# ══════════════════════════════════════════════════════════════════
h1('13.  Meeting Day Checklist')

h2('Before the Meeting — At Least One Week Prior')
for item in [
    'Log in to App Settings — change all six passwords from their defaults.',
    'Set organization name and election name in App Settings.',
    'Election Setup → Details: enter expected voters and absentee count.',
    'Election Setup → Nominees: enter nominee names, open positions, votes per voter. Save.',
    'Election Setup → Settings: configure completion message; confirm voter URL.',
    'Voter Tokens: generate tokens (one per eligible voter). Print and cut token cards.',
    'Test the voter URL on a phone — confirm the token entry screen appears.',
    'If running a motion vote: prepare the motion text and answer options.',
    'Print paper ballots.',
    'Brief the chairman on the results password and how to access the Election Dashboard.',
    'Brief the paper ballot volunteer on the paper ballot password and their screen.',
    'Store token cards securely until the meeting.',
]:
    p = doc.add_paragraph(); _space(p, 0, 4)
    p.add_run('☐   ' + item).font.size = Pt(11)

h2('Day of the Meeting')
for item in [
    'Confirm internet is working — test the voter URL on a phone.',
    'Log in to the admin hub on your device.',
    'Confirm all election setup is correct (nominees, voter count, voter URL).',
    'Open Election Dashboard on a separate device or browser tab for the chairman.',
    'Open Paper Ballot Entry on the volunteer\'s device.',
    'Distribute token cards at the entrance — one per eligible voter.',
    'Confirm paper ballot volunteer is logged in and the station is ready.',
]:
    p = doc.add_paragraph(); _space(p, 0, 4)
    p.add_run('☐   ' + item).font.size = Pt(11)

h2('After the Meeting')
for item in [
    'Chairman or election officer may screenshot Election Dashboard for the minutes.',
    'Use Reset All Election Data (Election Setup → Settings → Danger Zone) to clear the election.',
    'Use Reset & New Vote (Motion Vote Control) to clear motion vote data.',
    'Shred or store token cards securely.',
]:
    bullet(item)


# ══════════════════════════════════════════════════════════════════
# 14. SUGGESTED ORDER OF EVENTS
# ══════════════════════════════════════════════════════════════════
h1('14.  Suggested Order of Events')

body('Example: electing 2 board members, then a motion vote.')

h2('Board Election')
for step in [
    'Chairman introduces the nominees and explains the voting process.',
    'Chairman asks members to open the voter URL or scan the QR code on their token card.',
    'Election officer: Round Control → Start Round 1 Voting.',
    'Chairman: "Voting is now open. Please select your candidate(s) and submit."',
    'Paper ballot volunteer enters absentee ballots first (Absentee checkbox ticked), then remaining paper ballots.',
    'Election officer monitors Ballots In and Participation %. When satisfied: Close Voting → End Round → confirm.',
    'Election officer processes Round Transition: confirm elected candidates (Step 1); set advancing candidates if Round 2 is needed (Step 2); set votes per voter (Step 3).',
    'If all positions filled: Complete This Election. If not: Launch Next Round.',
    'Chairman announces results verbally (from Election Dashboard — vote counts not shown publicly).',
    'Voter phones display the election completion message.',
]:
    numbered(step)

h2('Motion Vote')
for step in [
    'Chairman introduces the motion and reads it to the membership.',
    'Election officer: Home Screen → Motion Vote → enter election password.',
    'Enter the motion text, confirm answer options and expected voters, set majority threshold.',
    'Click Open Motion Voting.',
    'Voter phones automatically switch to the motion ballot (members who voted in the election do not need to re-enter their token).',
    'Paper ballot volunteer enters paper motion votes (screen switches to motion mode automatically).',
    'Election officer: Close Voting → Mark Complete (confirm).',
    'Election officer announces the result verbally based on what is shown on the Motion Vote Control screen.',
]:
    numbered(step)

tip('Always run a complete test election before the actual meeting. Use at least two phones to confirm the full flow from token entry to election complete.')


# ══════════════════════════════════════════════════════════════════
# 15. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════
h1('15.  Troubleshooting')

table(
    ['Problem', 'Likely Cause', 'Solution'],
    [
        ['App password screen does not dismiss',
         'Wrong password entered.',
         'Re-enter the correct app password. If forgotten, contact your hosting administrator.'],
        ['Voter page shows "No Voting Configured"',
         'No election or motion vote is set up yet, or setup was saved with server offline.',
         'Complete Election Setup and save. Confirm the voter URL in Settings is correct.'],
        ['Voter cannot reach the voter page',
         'Wrong URL, or voter\'s phone has no internet connection.',
         'Confirm the URL on the token card matches your app URL. Confirm voter has internet access.'],
        ['"Failed to record vote" on voter\'s phone',
         'Phone lost internet connection momentarily.',
         'Have voter try again. If persistent, check that the app URL is reachable.'],
        ['"Invalid token" error',
         'Code mistyped or tokens were regenerated after printing.',
         'Double-check the 4-digit code. If tokens were regenerated, old cards are invalid — reprint.'],
        ['"Token already used" error',
         'Voter has already submitted a ballot this round, or a paper ballot was entered for their token.',
         'Voter has already voted this round. Verify with the paper ballot log if unsure.'],
        ['Voter phone stuck on "Voting Closed"',
         'Voting was closed in Round Control; voter was already on the ballot screen.',
         'Normal behaviour — tell voter to wait. Page updates automatically when next round opens.'],
        ['Vote counts not updating in Round Control',
         'Browser tab lost network connection.',
         'Refresh the browser tab. Counts update every 3 seconds when connected.'],
        ['Paper Ballot Entry form not appearing',
         'Voting is not currently open.',
         'Open voting in Round Control or Motion Vote Control. Screen updates within 4 seconds.'],
        ['Motion ballot not showing on phones after opening voting',
         'Voter phones have not yet polled for the update.',
         'Wait up to 3 seconds. If still not showing, ask voter to refresh their browser tab.'],
        ['Login fails on mobile (password rejected)',
         'Mobile keyboard autocapitalised or added a space.',
         'Type the password carefully. All password fields disable autocorrect. Do not copy-paste.'],
        ['Forgotten password',
         'Password was changed and not recorded.',
         'Contact your hosting administrator to reset passwords.'],
    ],
    col_widths=[4.0, 4.5, 7.5]
)


# ══════════════════════════════════════════════════════════════════
# 16. QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════
h1('16.  Quick Reference')

h2('Passwords — Fill In Before the Meeting')
table(
    ['Password', 'Default', 'Protects', 'Custom Password'],
    [
        ['App',          'votevote2024',   'Admin hub login (all officers)',             '____________________'],
        ['Admin',        'churchvoting',   'App Settings',                               '____________________'],
        ['Election',     'election2024',   'Election Setup, Round Control, Motion Vote', '____________________'],
        ['Results',      'results2024',    'Election Dashboard',                         '____________________'],
        ['Paper Ballot', 'paperentry2024', 'Paper Ballot Entry',                         '____________________'],
        ['Tokens',       'tokens2024',     'Voter Tokens',                               '____________________'],
    ],
    col_widths=[3.0, 3.0, 5.0, 5.0]
)

h2('App URL & Key Numbers')
table(
    ['', 'Value'],
    [
        ['Admin Hub URL',       '________________________________________________'],
        ['Voter Page URL',      '________________________________________________'],
        ['Expected Voters',     '________'],
        ['Absentee Votes',      '________'],
        ['Round 1 Majority',    'floor((Attending + Absentee) / 2) + 1  =  ________'],
        ['Round 2+ Majority',   'floor(Attending / 2) + 1  =  ________'],
        ['Motion Threshold',    '____________________________'],
    ],
    col_widths=[4.5, 11.5]
)

h2('Screen Navigation Summary')
table(
    ['Screen', 'Who', 'Password', 'Purpose'],
    [
        ['Home Screen',        'All officers',    'App password',   'Navigation hub.'],
        ['App Settings',       'Elect. Officer',  'Admin',          'Organization identity and all passwords.'],
        ['Election Hub',       'Elect. Officer',  '(none)',         'Entry point for election screens.'],
        ['Election Setup',     'Elect. Officer',  'Election',       'Configure nominees, voter count.'],
        ['Round Control',      'Elect. Officer',  'Election (hub)', 'Open/close voting, monitor, transitions.'],
        ['Voter Tokens',       'Token Admin',     'Tokens',         'Generate tokens, print cards.'],
        ['Motion Vote',        'Elect. Officer',  'Election',       'Configure and run motion vote.'],
        ['Paper Ballot Entry', 'Volunteer',       'Paper Ballot',   'Enter paper votes — election or motion.'],
        ['Election Dashboard', 'Chairman',        'Results',        'Live election status and vote counts.'],
        ['Voter Page',         'All voters',      'Token code',     'Submit digital ballots.'],
    ],
    col_widths=[3.5, 3.0, 3.2, 6.3]
)

# Footer
fp = doc.add_paragraph(); _space(fp, 24, 0); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('Board Voting App — Election Officer Manual  |  votingapp.ca\nFor Boards, Committees, and Organizations')
r.font.size = Pt(9); r.font.color.rgb = GREY

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'board_manual.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
