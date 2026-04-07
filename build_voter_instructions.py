#!/usr/bin/env python3
"""Generate voter_instructions.docx — 2 half-page instruction sheets per letter page."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x27, 0x44)
GOLD   = RGBColor(0xb8, 0x94, 0x2a)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREY   = RGBColor(0x66, 0x66, 0x66)
LGREY  = RGBColor(0x44, 0x44, 0x44)
GREEN  = RGBColor(0x2d, 0x6a, 0x4f)
LGREEN = RGBColor(0xe8, 0xf5, 0xee)
AMBER  = RGBColor(0x92, 0x60, 0x10)
LAMBER = RGBColor(0xff, 0xf8, 0xe6)

doc = Document()

# ── Letter page, tight margins ────────────────────────────────────────────────
for section in doc.sections:
    section.page_height = Inches(11)
    section.page_width  = Inches(8.5)
    section.top_margin    = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin   = Inches(0.5)
    section.right_margin  = Inches(0.5)

# ── XML helpers ───────────────────────────────────────────────────────────────
def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)

def set_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        if val:
            el.set(qn('w:val'),   val.get('val','single'))
            el.set(qn('w:sz'),    val.get('sz','4'))
            el.set(qn('w:space'),'0')
            el.set(qn('w:color'), val.get('color','auto'))
        else:
            el.set(qn('w:val'), 'none')
        tcB.append(el)
    tcPr.append(tcB)

def remove_table_borders(tbl):
    """Remove all outer and inner borders from a table."""
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tblEl.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBdr.append(el)
    tblPr.append(tblBdr)

def set_row_height(row, height_inches, exact=True):
    tr    = row._tr
    trPr  = tr.get_or_add_trPr()
    trH   = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_inches * 1440)))
    trH.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    trPr.append(trH)

def para_space(p, before=0, after=0):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(14)

# ── Step circle helper ────────────────────────────────────────────────────────
def add_step(cell_or_doc, number, heading, detail, number_colour='1A2744', is_last=False):
    """Add a numbered step row (circle number | heading | detail) inside a 3-col inner table."""
    inner = cell_or_doc.add_table(rows=1, cols=3)
    remove_table_borders(inner)
    inner.alignment = WD_TABLE_ALIGNMENT.LEFT
    inner.allow_autofit = False

    # Column widths: circle | heading+detail | spacer
    widths = [Inches(0.55), Inches(5.85), Inches(1.1)]
    for ci, w in enumerate(widths):
        for row in inner.rows:
            row.cells[ci].width = w

    # Circle cell
    circ_cell = inner.rows[0].cells[0]
    circ_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade_cell(circ_cell, number_colour)
    cp = circ_cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(cp, before=3, after=3)
    cr = cp.add_run(str(number))
    cr.bold = True
    cr.font.size = Pt(14)
    cr.font.color.rgb = WHITE

    # Text cell
    text_cell = inner.rows[0].cells[1]
    text_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    hp = text_cell.paragraphs[0]
    para_space(hp, before=2, after=1)
    hr = hp.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(11.5)
    hr.font.color.rgb = NAVY

    dp = text_cell.add_paragraph()
    para_space(dp, before=0, after=4 if not is_last else 2)
    dr = dp.add_run(detail)
    dr.font.size = Pt(10)
    dr.font.color.rgb = LGREY

    # Spacer cell (empty)
    spacer_cell = inner.rows[0].cells[2]
    set_borders(spacer_cell)

    return inner


# ════════════════════════════════════════════════════════════════════════════════
# Build one instruction card inside a given table cell
# ════════════════════════════════════════════════════════════════════════════════
def build_card(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade_cell(cell, 'FAFAF8')

    # ── Header bar ─────────────────────────────────────────────────────────────
    hdr_tbl = cell.add_table(rows=1, cols=1)
    remove_table_borders(hdr_tbl)
    hdr_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cell = hdr_tbl.rows[0].cells[0]
    shade_cell(hdr_cell, '1A2744')
    set_row_height(hdr_tbl.rows[0], 0.52)
    hp = hdr_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(hp, before=6, after=4)
    htitle = hp.add_run('HOW TO CAST YOUR VOTE')
    htitle.bold = True
    htitle.font.size = Pt(15)
    htitle.font.color.rgb = WHITE
    hsub = hdr_cell.add_paragraph()
    hsub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(hsub, before=0, after=6)
    hsubr = hsub.add_run('Follow these steps on your mobile device')
    hsubr.font.size = Pt(9)
    hsubr.font.color.rgb = RGBColor(0xb8, 0xc8, 0xe0)

    # ── Gold divider line ──────────────────────────────────────────────────────
    div_tbl = cell.add_table(rows=1, cols=1)
    remove_table_borders(div_tbl)
    div_cell = div_tbl.rows[0].cells[0]
    shade_cell(div_cell, 'B8942A')
    set_row_height(div_tbl.rows[0], 0.05, exact=True)
    div_cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    div_cell.paragraphs[0].paragraph_format.space_after  = Pt(0)

    # ── Steps ──────────────────────────────────────────────────────────────────
    steps = [
        (1, 'Open the Voter Page',
         'Scan the QR code on your token card with your phone camera,\n'
         'or type the voter URL into your mobile browser.'),
        (2, 'Enter Your 4-Digit Token Code',
         'Type the code printed on your token card and tap  Submit Token.\n'
         'Each voter has a unique code — do not share it with others.'),
        (3, 'Select Your Candidate(s)',
         'Tap a candidate\'s name to select them. A coloured highlight and checkmark\n'
         'confirm your choice. You may select up to the number shown at the top.'),
        (4, 'Submit Your Vote',
         'Tap  Submit My Vote  when you are done. A confirmation screen appears.\n'
         'Your vote is final — it cannot be changed after submission.'),
        (5, 'Wait for the Next Round (if applicable)',
         'Keep the page open. If another round or office opens, a notification\n'
         'appears automatically — tap  Vote Now  to continue with the same token card.',
         ),
    ]

    for i, step in enumerate(steps):
        num, heading, detail = step
        is_last = (i == len(steps) - 1)
        add_step(cell, num, heading, detail, number_colour='1A2744', is_last=is_last)

    # ── Tips box ───────────────────────────────────────────────────────────────
    tip_tbl = cell.add_table(rows=1, cols=1)
    remove_table_borders(tip_tbl)
    tip_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tip_cell = tip_tbl.rows[0].cells[0]
    shade_cell(tip_cell, 'FFF8E6')
    set_borders(tip_cell,
        top    ={'val':'single','sz':'6','color':'B8942A'},
        bottom ={'val':'single','sz':'6','color':'B8942A'},
        left   ={'val':'single','sz':'6','color':'B8942A'},
        right  ={'val':'single','sz':'6','color':'B8942A'},
    )

    tp = tip_cell.paragraphs[0]
    para_space(tp, before=4, after=2)
    tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tlabel = tp.add_run('HELPFUL TIPS     ')
    tlabel.bold = True
    tlabel.font.size = Pt(9)
    tlabel.font.color.rgb = AMBER

    tips = [
        '"Voting Round Closed" — the chairman has paused voting. Keep the page open; it will update automatically.',
        '"Token already used" — you have already voted this round. One vote per round per office.',
        'Your vote is completely anonymous. No one can see what you selected.',
        'One token card covers all rounds and both offices (Elder and Deacon).',
        'If you have difficulty, ask the paper ballot volunteer for a paper ballot instead.',
    ]
    for tip in tips:
        tipp = tip_cell.add_paragraph()
        para_space(tipp, before=0, after=2)
        tipr = tipp.add_run('▸  ' + tip)
        tipr.font.size = Pt(9)
        tipr.font.color.rgb = LGREY

    # bottom padding in tip cell
    lastp = tip_cell.add_paragraph()
    para_space(lastp, before=0, after=3)

    # ── Footer ─────────────────────────────────────────────────────────────────
    fp = cell.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(fp, before=4, after=2)
    fr = fp.add_run('Your vote is anonymous  ·  Keep this page open until the election is complete')
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY
    fr.italic = True


# ════════════════════════════════════════════════════════════════════════════════
# Page layout — 2-row outer table, one card per row, dashed cut line between
# ════════════════════════════════════════════════════════════════════════════════
outer = doc.add_table(rows=3, cols=1)
remove_table_borders(outer)
outer.alignment = WD_TABLE_ALIGNMENT.LEFT
outer.allow_autofit = False

# Row 0 — Card 1
set_row_height(outer.rows[0], 4.85, exact=False)
build_card(outer.rows[0].cells[0])

# Row 1 — Cut line
set_row_height(outer.rows[1], 0.25, exact=True)
cut_cell = outer.rows[1].cells[0]
shade_cell(cut_cell, 'FFFFFF')
set_borders(cut_cell,
    top   ={'val':'dashed','sz':'6','color':'AAAAAA'},
    bottom={'val':'none'},
    left  ={'val':'none'},
    right ={'val':'none'},
)
cp2 = cut_cell.paragraphs[0]
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(cp2, before=2, after=0)
cr2 = cp2.add_run('✂   fold or cut here   ✂')
cr2.font.size = Pt(8)
cr2.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
cr2.italic = True

# Row 2 — Card 2 (identical)
set_row_height(outer.rows[2], 4.85, exact=False)
build_card(outer.rows[2].cells[0])

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/Users/johan/Documents/_CLAUDE CODE/Church Election App/voter_instructions.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
