#!/usr/bin/env python3
"""
Church Voting App — Election Officer Manual
Generates church_manual.docx   (run: python3 build_church_manual.py)
Hosted solution edition — for use with votingapp.ca managed hosting.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

NAVY   = RGBColor(0x1a, 0x27, 0x44)
GOLD   = RGBColor(0xb8, 0x94, 0x2a)
GREEN  = RGBColor(0x2d, 0x6a, 0x4f)
RED    = RGBColor(0x9b, 0x23, 0x35)
GREY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xff, 0xff, 0xff)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

def _space(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def shade_cell(cell, hex_colour):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph(); _space(p, 18, 6)
    p.paragraph_format.page_break_before = True
    r = p.add_run(text); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '1a2744')
    pBdr.append(bot); pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph(); _space(p, 12, 4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = NAVY

def h3(text):
    p = doc.add_paragraph(); _space(p, 8, 3)
    r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GREY

def body(text):
    p = doc.add_paragraph(); _space(p, 0, 5)
    r = p.add_run(text); r.font.size = Pt(11)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet'); _space(p, 0, 3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.add_run(text).font.size = Pt(11)

def numbered(text):
    p = doc.add_paragraph(style='List Number'); _space(p, 0, 3)
    p.paragraph_format.left_indent = Cm(0.5)
    p.add_run(text).font.size = Pt(11)

def callout(label, text, bg_hex, border_hex):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0); shade_cell(cell, bg_hex)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '12' if side in ('top', 'left') else '4')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), border_hex)
        tcBorders.append(el)
    tcPr.append(tcBorders)
    lp = cell.paragraphs[0]; _space(lp, 2, 2)
    lr = lp.add_run(f'{label}  '); lr.bold = True; lr.font.size = Pt(10)
    lp.add_run(text).font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def note(text):    callout('NOTE:',            text, 'E8F0FA', '1A3A5C')
def tip(text):     callout('TIP:',             text, 'F0FAF0', '2D6A4F')
def warning(text): callout('IMPORTANT:',       text, 'FFF4E0', 'B8942A')
def danger(text):  callout('ACTION REQUIRED:', text, 'FDF0F0', '9B2335')

def table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]; shade_cell(cell, '1A2744')
        p = cell.paragraphs[0]; _space(p, 2, 2)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE
    for ri, row_data in enumerate(rows):
        bg = 'F8F6F1' if ri % 2 else 'FFFFFF'
        for ci, text in enumerate(row_data):
            cell = tbl.rows[ri + 1].cells[ci]; shade_cell(cell, bg)
            p = cell.paragraphs[0]; _space(p, 2, 2)
            p.add_run(str(text)).font.size = Pt(10)
    if col_widths:
        for row in tbl.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
cover = doc.add_paragraph(); cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
_space(cover, 80, 6)
r = cover.add_run('Church Voting App'); r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY

sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; _space(sub, 6, 10)
r = sub.add_run('Election Officer Manual'); r.font.size = Pt(16); r.font.color.rgb = GOLD

ver = doc.add_paragraph(); ver.alignment = WD_ALIGN_PARAGRAPH.CENTER; _space(ver, 0, 30)
r = ver.add_run('Hosted Solution Edition — votingapp.ca'); r.font.size = Pt(11); r.font.color.rgb = GREY

for line in [
    'Office Bearer Elections · Congregational Voting',
    'For Reformed Church Congregational Meetings',
    '',
    'Confidential — For Election Officers & Chairman Use Only',
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _space(p, 0, 4)
    r = p.add_run(line); r.font.size = Pt(11)
    r.bold = 'Confidential' in line
    r.font.color.rgb = NAVY if 'Confidential' in line else GREY

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
toc = doc.add_paragraph(); _space(toc, 0, 8)
r = toc.add_run('Table of Contents'); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

toc_entries = [
    ('1.',     'System Overview'),
    ('2.',     'Accessing the App'),
    ('3.',     'App Settings'),
    ('   3a.', 'Church Identity'),
    ('   3b.', 'Passwords'),
    ('4.',     'Home Screen'),
    ('5.',     'Election Setup'),
    ('   5a.', 'Congregation Details'),
    ('   5b.', 'Elder Election'),
    ('   5c.', 'Deacon Election'),
    ('   5d.', 'Settings — Voter URL & QR Code'),
    ('6.',     'Voter Tokens'),
    ('7.',     'Congregational Vote Setup'),
    ('8.',     'Running the Election — Round Control'),
    ('   8a.', 'Starting a Round'),
    ('   8b.', 'Opening & Closing Voting'),
    ('   8c.', 'Monitoring Participation'),
    ('   8d.', 'Round Transition'),
    ('   8e.', 'Completing an Office'),
    ('9.',     'Paper Ballot Entry'),
    ('10.',    'The Voter Page'),
    ('11.',    'Election Dashboard'),
    ('12.',    'Congregational Vote — Running the Vote'),
    ('   12a.','Voting Control'),
    ('   12b.','Voting Dashboard'),
    ('   12c.','Paper Ballot Entry — Congregational Vote'),
    ('13.',    'Roles & Responsibilities'),
    ('14.',    'Meeting Day Checklist'),
    ('15.',    'Suggested Order of Events'),
    ('16.',    'Troubleshooting'),
    ('17.',    'Quick Reference'),
]
for num, title in toc_entries:
    p = doc.add_paragraph(); _space(p, 0, 3)
    indent = len(num) - len(num.lstrip())
    p.paragraph_format.left_indent = Cm(indent * 0.2)
    top = not num.startswith('   ')
    r1 = p.add_run(num + '  '); r1.font.size = Pt(11); r1.bold = top
    r1.font.color.rgb = NAVY if top else GREY
    r2 = p.add_run(title); r2.font.size = Pt(11)
    r2.font.color.rgb = NAVY if top else GREY

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ══════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════
h1('1.  System Overview')

body(
    'The Church Voting App is a browser-based voting platform for Reformed church congregational '
    'meetings. Your church\'s instance is hosted and managed on votingapp.ca — there is no '
    'software to install and no server to run. You access the admin hub through your web browser, '
    'and voters use the ballot page on their own phones or tablets.'
)

h2('What the App Supports')
for item in [
    'Multi-round office bearer elections — Elder and Deacon, independently configured',
    'Congregational vote — motions and resolutions with configurable majority threshold',
    'Digital ballots via personal phones using unique 4-digit token codes',
    'Paper ballot entry for members who prefer to vote on paper',
    'Absentee vote support for both the election and the congregational vote',
    'Real-time vote count updates — Round Control and dashboards refresh every 3 seconds',
    'One token card per voter covers all rounds, both offices, and the congregational vote',
    'Password-protected access levels so each role sees only what they need',
]:
    bullet(item)

h2('The Three Pages')
table(
    ['Page', 'Who Uses It', 'Purpose'],
    [
        ['Admin Hub\n(your app URL/)', 'Election officer, chairman, volunteers', 'All setup, round control, results, token printing'],
        ['Voter Page\n(your app URL/vote.html)', 'All eligible voters', 'Enter token, submit election ballot, cast congregational vote'],
        ['App Settings\n(accessible from Admin Hub)', 'Election officer only', 'Church identity, colours, logo, and all passwords'],
    ],
    col_widths=[4.5, 4.5, 7.0]
)

note('Your specific app URL (e.g. https://firstchurch.votingapp.ca) is provided by your hosting administrator. All pages are accessed via that same base URL.')


# ══════════════════════════════════════════════════════════════════
# 2. ACCESSING THE APP
# ══════════════════════════════════════════════════════════════════
h1('2.  Accessing the App')

body(
    'The Church Voting App runs entirely in a web browser. No installation is required on any device.'
)

h2('Admin Hub')
numbered('Open your web browser (Chrome, Safari, Edge, or Firefox).')
numbered('Navigate to your app URL — e.g. https://firstchurch.votingapp.ca')
numbered('Enter the app password on the login screen and tap Log In.')
numbered('The Home Screen appears, showing the four main navigation tiles.')

h2('Voter Devices')
numbered('Voters open the voter page in any browser on their phone or tablet.')
numbered('They can type the voter URL directly, or scan the QR code printed on their token card.')
numbered('No app installation, no account, no login — just the token code.')

note('The voter URL is your app URL with /vote.html at the end (e.g. https://firstchurch.votingapp.ca/vote.html). It is shown in Election Setup → Settings and printed on every token card.')

warning(
    'The app requires an active internet connection on all devices. Voters must be able to reach '
    'your app URL from their phones. Confirm this works before the meeting by testing on a phone.'
)

h2('Device Requirements')
table(
    ['Device', 'Requirements'],
    [
        ['Admin device (election officer)', 'Any modern browser. Keep this tab open throughout the meeting.'],
        ['Chairman device', 'Any browser — used for Election Dashboard and Voting Dashboard only.'],
        ['Paper ballot station', 'Any browser — used for Paper Ballot Entry only. Dedicate a separate device.'],
        ['Voter phones/tablets', 'Any browser. Mobile data or Wi-Fi — either works if it can reach the app URL.'],
    ],
    col_widths=[5.0, 11.0]
)


# ══════════════════════════════════════════════════════════════════
# 3. APP SETTINGS
# ══════════════════════════════════════════════════════════════════
h1('3.  App Settings')

body(
    'App Settings is where you configure your church\'s identity and manage all passwords. '
    'Access it from the Home Screen by clicking the ⚙ Settings link at the bottom of the page. '
    'The admin password is required.'
)

h2('3a.  Church Identity')
body('Fill in the following fields so the app reflects your congregation:')
table(
    ['Field', 'Description', 'Example'],
    [
        ['Congregation Name', 'Displayed on the voter page header and all token cards.', 'First Reformed Church'],
        ['Denomination', 'Select from the dropdown — affects the logo shown on the login screen.', 'Canadian Reformed'],
        ['Logo', 'Upload your church logo (optional). Appears on the login screen.', '(upload image)'],
        ['Primary Colour', 'Main accent colour — used for buttons, badges, and headers.', '#1a2744'],
        ['Secondary Colour', 'Accent colour for highlights and secondary elements.', '#b8942a'],
        ['Heading Font', 'Choose the typeface used for headings throughout the app.', 'Lato'],
    ],
    col_widths=[4.0, 8.0, 4.0]
)
body('Click Save Identity Settings to apply all identity changes.')
tip('Set your congregation name before printing token cards — the name appears on every card.')

h2('3b.  Passwords')
body(
    'The app uses six independent passwords so each role only accesses what they need. '
    'Change all passwords from their defaults before the meeting.'
)
table(
    ['Password', 'Default', 'Protects', 'Held By'],
    [
        ['Admin Password',       'election2024',   'App Settings',                                         'Election Officer only'],
        ['Election Password',    'election2024',   'Election Setup and Round Control',                     'Election Officer only'],
        ['Voting Password',      'voting2024',     'Voting Setup and Voting Control',                      'Election Officer only'],
        ['Results Password',     'results2024',    'Election Dashboard and Voting Dashboard',              'Chairman only'],
        ['Tokens Password',      'tokens2024',     'Voter Tokens section',                                 'Token administrator'],
        ['Paper Ballot Password','paperentry2024', 'Paper Ballot Entry station',                           'Paper ballot volunteer'],
    ],
    col_widths=[3.8, 3.0, 5.2, 4.0]
)
danger(
    'Change ALL passwords before the meeting. The defaults above are publicly documented. '
    'Go to App Settings to change each one. Record your chosen passwords securely and '
    'distribute each password only to the person who needs it.'
)
body(
    'To change a password: enter the current password in the "Current Password" field, '
    'type the new password twice, and click the Update button for that password. '
    'If you are still using the default, leave the current password field blank.'
)
note(
    'All password fields have autocapitalize, autocorrect, and spellcheck disabled to prevent '
    'mobile keyboard interference. Type carefully — do not copy-paste from other apps.'
)


# ══════════════════════════════════════════════════════════════════
# 4. HOME SCREEN
# ══════════════════════════════════════════════════════════════════
h1('4.  Home Screen')

body(
    'After logging in with the app password, the Home Screen shows four navigation tiles '
    'and the Settings link. A yellow warning banner appears at the top if the default app '
    'password is still in use.'
)

table(
    ['Tile / Link', 'Destination', 'Password Required'],
    [
        ['Office Bearer Election', 'Election Hub — Setup, Round Control, Dashboard',        'Election password (at hub)'],
        ['Congregational Vote',    'Voting Hub — Voting Setup, Control, Dashboard',         'Voting password (at hub)'],
        ['Paper Ballot Entry',     'Unified paper ballot station',                           'Paper ballot password'],
        ['Voter Tokens',           'Token generation, printing, and QR code',               'Tokens password'],
        ['⚙ Settings',            'App Settings — identity and all passwords',              'Admin password'],
    ],
    col_widths=[4.2, 6.8, 5.0]
)

tip('The election officer tab should stay open on the Office Bearer Election hub throughout the meeting. Use separate browser tabs for the Election Dashboard (chairman) and Paper Ballot Entry (volunteer).')


# ══════════════════════════════════════════════════════════════════
# 5. ELECTION SETUP
# ══════════════════════════════════════════════════════════════════
h1('5.  Election Setup')

body(
    'Election Setup is used before the meeting to configure the office bearer election. '
    'From the Home Screen, click Office Bearer Election, then click Election Setup. '
    'The election password is required.'
)
body('Election Setup has four tabs: Details, Elder Election, Deacon Election, and Settings.')
note('Configure only the offices being held. If only Elders are being elected this year, leave the Deacon tab empty — the system skips unconfigured offices automatically.')

h2('5a.  Congregation Details')
body('Click the Details tab and fill in:')
table(
    ['Field', 'Description', 'Example'],
    [
        ['Congregation Name',           'Full name shown on token cards, voter page, and dashboard.', 'First Reformed Church'],
        ['Meeting Date',                'Date of the congregational meeting.',                         '2026-11-15'],
        ['Expected Voters (attending)', 'Eligible male voters attending in person. Used for majority calculation and participation %.', '80'],
        ['Absentee Votes (Round 1)',    'Pre-meeting paper ballots received. Added to the majority base in Round 1 only.', '5'],
    ],
    col_widths=[4.5, 7.5, 4.0]
)
body(
    'Below these fields the system shows the Majority Required threshold. '
    'Majority = floor(n/2) + 1. Round 1 uses Expected Voters + Absentee Votes. '
    'Round 2 and later use Expected Voters only.'
)
tip('Click Save Details after entering these values, or the system saves automatically when you move between fields.')
warning('Enter the absentee count before the meeting. Each absentee ballot is entered individually during Round 1 via Paper Ballot Entry, with the Absentee checkbox ticked.')

h2('5b.  Elder Election Configuration')
body('Click the Elder Election tab:')
table(
    ['Field', 'Description'],
    [
        ['Nominees',             'One name per line. All men nominated for the office of Elder.'],
        ['Positions to Fill',    'Number of Elder positions open in this election.'],
        ['Votes per Voter (R1)', 'How many candidates each voter may select in Round 1. Typically equals the number of positions.'],
    ],
    col_widths=[4.5, 11.5]
)
body('Click Save Elder Setup when done.')
warning('Saving setup resets all in-progress election data for that office. Never click Save during an active voting round.')

h2('5c.  Deacon Election Configuration')
body(
    'Click the Deacon Election tab. Enter nominees, positions, and votes per voter — '
    'same process as Elder. Click Save Deacon Setup. '
    'Leave this tab empty if no Deacon election is being held this meeting.'
)

h3('Printing Paper Ballots')
bullet('Enter the number of Elder ballots needed and click Print Elder Ballots.')
bullet('Enter the number of Deacon ballots and click Print Deacon Ballots.')
bullet('Each ballot shows the office, congregation name, positions to fill, and nominee names with checkboxes.')
note('The same printed ballots can be reused in all rounds. The chairman verbally announces which candidates are on the ballot each round.')

h2('5d.  Settings — Voter URL & QR Code')
body(
    'The Settings tab under Election Setup manages the voter page URL, QR code, and password changes '
    'specific to the election. The voter URL is your app URL with /vote.html appended.'
)

h3('Voter Page URL')
body(
    'This field should already be set to your correct voter URL if your hosting is configured. '
    'Verify it shows the correct address. Click Save URL if you make any changes. '
    'The QR code updates automatically.'
)
warning('Confirm the voter URL before printing token cards. The URL and QR code printed on every card come from this saved value.')

h3('Reset All Election Data')
body(
    'The red Reset All Election Data button (Danger Zone) permanently erases all nominees, '
    'ballots, tokens, results, and resets passwords to defaults. Two confirmation prompts prevent '
    'accidental activation.'
)
danger('Reset All Election Data cannot be undone. Use only to prepare for a fresh election — never during a meeting.')


# ══════════════════════════════════════════════════════════════════
# 6. VOTER TOKENS
# ══════════════════════════════════════════════════════════════════
h1('6.  Voter Tokens')

body(
    'Each eligible voter receives one token card before the meeting. The 4-digit code on '
    'the card is their key to vote in all election rounds, both offices, and the congregational '
    'vote — they never need a new card.'
)
body('Access from the Home Screen: click Voter Tokens and enter the tokens password.')

h2('Generating Tokens')
numbered('Enter the total number of eligible voters (e.g. 85 — include both in-person and absentee).')
numbered('Click Generate Tokens. The system creates unique 4-digit codes.')
numbered('The token grid displays all codes. Used tokens are shown with a grey badge.')
warning('Generate tokens only once before the meeting. Regenerating during a live election invalidates all distributed cards.')

h2('Printing Token Cards')
body(
    'Click Print Token Cards to open the print view. Each card shows: congregation name, '
    'the 4-digit code in large text, a QR code linking to the voter page, and the voter URL. '
    'Print on card stock and cut into individual cards. Distribute one per eligible voter at the entrance.'
)
tip('Print and cut cards the day before the meeting. Use A4 card stock for best results. Have a door steward distribute cards only to verified eligible voters.')

h2('Voter URL & QR Code Display')
body(
    'The Voter Tokens screen also displays the current voter URL and a large QR code. '
    'Scan the QR code to confirm it opens the correct voter page before the meeting.'
)

h2('Changing the Tokens Password')
body('Enter the current tokens password, type the new password twice, and click Update Tokens Password.')


# ══════════════════════════════════════════════════════════════════
# 7. CONGREGATIONAL VOTE SETUP
# ══════════════════════════════════════════════════════════════════
h1('7.  Congregational Vote Setup')

body(
    'The Congregational Vote handles motions and resolutions. It is independent of the office '
    'bearer election and can be run before, after, or without an election.'
)
body(
    'From the Home Screen, click Congregational Vote and enter the voting password to reach '
    'the Voting Hub. From there, click Voting Setup.'
)

table(
    ['Field', 'Description', 'Example'],
    [
        ['Vote Question',      'The motion or resolution text members will vote on.',             'Do you support the proposed building fund resolution?'],
        ['Answer Options',     'The choices members select. At least two required. Add or remove as needed.', 'In favour / Not in favour'],
        ['Expected Voters',    'Eligible voters for this vote. Defaults to the election value if left at 0.', '80'],
        ['Absentee Voters',    'Number of absentee ballots to include in the majority calculation.', '5'],
        ['Majority Threshold', '50%+1 (simple majority), 75% of eligible voters, or a custom percentage.', '50%+1'],
    ],
    col_widths=[4.0, 8.0, 4.0]
)
body('Click Save Voting Setup. A green confirmation appears.')
warning('Saving Voting Setup resets all votes for the current vote. Never save while voting is open or after votes have been cast.')
tip('Once the vote is configured, a "▶ Voting Control" shortcut button appears in the Voting Hub header for quick access during the meeting.')


# ══════════════════════════════════════════════════════════════════
# 8. RUNNING THE ELECTION — ROUND CONTROL
# ══════════════════════════════════════════════════════════════════
h1('8.  Running the Election — Round Control')

body(
    'Round Control is used during the meeting to manage the office bearer election. '
    'From the Office Bearer Election hub, click Round Control. The election password is required.'
)
body(
    'The header changes colour by office (warm brown for Elder, deep blue for Deacon). '
    'Four stat boxes update live every 3 seconds. The control bar has three buttons.'
)

h2('8a.  Starting a Round')
body(
    'When no election is in progress, Round Control shows a "Ready to Start" panel for '
    'the first configured office, listing nominees and expected voters. '
    'Click Start Round 1 Voting to begin.'
)
tip('Do not click Start Round 1 until the chairman has introduced the candidates. Once clicked, the ballot is immediately available on voters\' phones.')

h2('8b.  Opening & Closing Voting')
table(
    ['Button', 'Action', 'When to Use'],
    [
        ['▶ Open Voting',  'Opens voting — members can now submit ballots.', 'At the start of each voting period, after the chairman\'s announcement.'],
        ['■ Close Voting', 'Closes voting — no further ballots accepted.', 'When participation is sufficient, before processing results.'],
        ['📊 End Round',   'Confirms the round and opens the Round Transition screen.', 'After voting is closed. Requires voting to be closed first.'],
    ],
    col_widths=[3.5, 6.0, 6.5]
)
warning('End Round requires voting to be closed first. A confirmation dialog shows the ballot count — click Confirm to proceed.')
body('When voting is closed, all voter screens update within 3 seconds to show "Voting Round Closed — please wait for further instructions."')

h2('8c.  Monitoring Participation')
body('The stat boxes update every 3 seconds:')
table(
    ['Statistic', 'Meaning'],
    [
        ['Ballots In',      'Total ballots this round (digital + paper combined).'],
        ['Paper Ballots',   'Ballots entered via the Paper Ballot Entry station.'],
        ['Absentee',        'Paper ballots marked as absentee. Shown in Round 1 only.'],
        ['Expected Voters', 'Round 1: attending + absentee. Round 2+: attending only.'],
        ['Participation %', 'Ballots In ÷ Expected Voters × 100. Shown green when ≥ 50%.'],
    ],
    col_widths=[4.5, 11.5]
)
note('Live vote counts per candidate are visible only on Round Control and the Election Dashboard — not on any public screen or voter phone.')

h2('8d.  Round Transition')
body('After clicking End Round and confirming, the Round Transition screen opens with four sections:')
for step in [
    'Final Results — all candidates ranked by votes received.',
    'Step 1 — Mark as Elected: candidates who reached majority are pre-checked. Adjust checkboxes if needed.',
    'Step 2 — Select Who Advances: choose which remaining candidates go to the next round (all advance by default). Uncheck to exclude.',
    'Step 3 — Next Round Settings: set votes per voter for the next round. Auto-filled with remaining positions.',
    'If all positions are filled: a green banner appears. Steps 2 and 3 are hidden. Only Complete This Office is available.',
    'Click Launch Next Round (opens Round Control with voting closed so the chairman can announce first) OR Complete This Office.',
]:
    numbered(step)

h2('8e.  Completing an Office')
body('When all positions are filled, click Complete This Office. The system:')
bullet('Records elected candidates for the office.')
bullet('If the other office is configured and not started: shows "Ready to Start" for that office.')
bullet('If both offices are complete: shows the Election Complete screen.')

body('All voter phones simultaneously display the thank-you message once the election is complete.')


# ══════════════════════════════════════════════════════════════════
# 9. PAPER BALLOT ENTRY
# ══════════════════════════════════════════════════════════════════
h1('9.  Paper Ballot Entry')

body(
    'The Paper Ballot Entry screen is used by a dedicated volunteer to enter votes for '
    'members who voted on paper. It is accessed from the Home Screen — the paper ballot '
    'password is required. The screen automatically detects whether an election or '
    'congregational vote is open and shows the correct form.'
)
warning('Open Paper Ballot Entry on a dedicated device or a separate browser tab. It is operated by a volunteer, not the election officer.')

h2('Entering an Election Ballot')
numbered('Confirm the correct office and round are shown at the top of the screen.')
numbered('Click each candidate\'s name from the paper ballot. Selected candidates are highlighted with a checkmark.')
numbered('(Round 1 only) If the ballot is from an absentee voter, tick the "Absentee ballot" checkbox before submitting.')
numbered('Click Submit Paper Ballot. The vote is recorded and the form resets.')
numbered('The submitted ballot appears in the log at the bottom of the screen.')

h2('Editing or Deleting a Ballot')
bullet('To edit: click the ✏ Edit button next to a ballot in the log. Make corrections and click Update Paper Ballot.')
bullet('To delete the most recent entry: click Delete Last Entry (red button) and confirm.')
note('Only the most recently submitted ballot can be deleted. Use the Edit button to correct earlier entries.')

h2('When No Voting Is Open')
body(
    'If no election or congregational vote is currently open, the screen shows '
    '"Voting is currently closed." The form appears automatically within 4 seconds '
    'when voting opens.'
)


# ══════════════════════════════════════════════════════════════════
# 10. THE VOTER PAGE
# ══════════════════════════════════════════════════════════════════
h1('10.  The Voter Page')

body(
    'Voters use a single URL throughout the entire meeting — for both the office bearer '
    'election and the congregational vote. The page automatically detects what is active '
    'and shows the correct ballot. No navigation or page refreshing is needed.'
)

h2('Voting Flow — Office Bearer Election')
numbered('Voter opens the voter URL or scans the QR code on their token card.')
numbered('Voter enters their 4-digit token code and taps Submit Token.')
numbered('The ballot appears: office name, round number, candidates, and how many to select.')
numbered('Voter taps each candidate to select. A coloured highlight confirms the choice.')
numbered('Voter taps Submit My Vote. A "Vote Submitted!" confirmation appears.')
numbered('If another round or office opens, a notification appears — voter taps Vote Now.')
note('The same token card is used for all rounds and both offices. Voters never need to re-enter their token for a new round of the same office.')

h2('Voting Flow — Congregational Vote')
numbered('If the voter already voted in the election, the page transitions automatically — no re-entry of the token is needed.')
numbered('If no election was held, the voter enters their token as usual.')
numbered('The congregational vote ballot appears with the vote question and answer options.')
numbered('Voter selects their answer and taps Cast My Vote.')
numbered('A "Vote Recorded" confirmation shows the question and selected answer.')

h2('All Voter Screen States')
table(
    ['State', 'When It Appears'],
    [
        ['Token Entry',           'Voting is open; voter has not yet entered their token this session.'],
        ['Election Ballot',       'Token entered; office bearer election is open.'],
        ['Congregational Ballot', 'Token entered; congregational vote is open.'],
        ['Vote Submitted',        'Election ballot confirmed. Page polls for next round or office.'],
        ['Vote Recorded',         'Congregational vote confirmed.'],
        ['Voting Closed',         'Round or vote closed while voter is on ballot screen.'],
        ['Waiting',               'Voting not yet open or between rounds.'],
        ['Election Complete',     'All offices done — thank-you message shown.'],
        ['Vote Complete',         'Congregational vote marked complete.'],
    ],
    col_widths=[5.0, 11.0]
)
warning('Votes are final once submitted. Digital ballots cannot be changed. Only paper ballots can be edited via the Paper Ballot Entry station.')


# ══════════════════════════════════════════════════════════════════
# 11. ELECTION DASHBOARD
# ══════════════════════════════════════════════════════════════════
h1('11.  Election Dashboard')

body(
    'The Election Dashboard is a read-only view for the chairman showing live vote counts '
    'and election status, refreshing every 3 seconds. '
    'From the Office Bearer Election hub, click Election Dashboard and enter the results password.'
)

h2('What the Dashboard Shows')
table(
    ['Section', 'Content'],
    [
        ['Status Card',       'Current office and round, voting status pill (Open/Closed/Configuring), majority required, votes per voter, candidates on ballot.'],
        ['Info Strip',        'Expected Voters and Absentee Votes pills.'],
        ['Results Section',   'All candidates with proportional vote bars, vote counts, and ★ Majority / ✓ Elected badges.'],
    ],
    col_widths=[4.5, 11.5]
)

body('The results section adapts to the election state:')
bullet('Before Round 1: shows nominees, positions, and Round 1 majority threshold.')
bullet('During a round: shows live vote counts and bars labelled "● Voting Open".')
bullet('After End Round (before transition): shows a "Round N Results — Awaiting transition" card.')
bullet('After transition: shows committed historical results for all completed rounds.')
tip('The "Awaiting transition" card lets the chairman review results and prepare their announcement before the election officer processes the transition.')
warning('The dashboard shows actual vote counts. Keep the results password confidential and never leave the dashboard visible on an unattended screen.')


# ══════════════════════════════════════════════════════════════════
# 12. CONGREGATIONAL VOTE — RUNNING THE VOTE
# ══════════════════════════════════════════════════════════════════
h1('12.  Congregational Vote — Running the Vote')

body(
    'Once configured in Voting Setup, the congregational vote is managed through Voting Control. '
    'The chairman monitors results on the Voting Dashboard. Both are accessed from the '
    'Voting Hub (Congregational Vote → enter voting password).'
)

h2('12a.  Voting Control')
body('Managing the congregational vote:')
table(
    ['Button', 'Action', 'Notes'],
    [
        ['▶ Open Voting',      'Opens voting — members can now cast their vote.', 'Voter phones switch to the congregational vote ballot automatically.'],
        ['■ Close Voting',     'Closes voting.', 'Voter phones update within 3 seconds.'],
        ['✅ Mark Complete',   'Finalises the vote permanently.', 'Disabled while voting is open. Requires confirmation.'],
        ['🔄 Reset & New Vote','Clears all votes and returns to unconfigured state.', 'Appears only after vote is marked complete.'],
    ],
    col_widths=[3.8, 5.7, 6.5]
)

body('The stat boxes update every 3 seconds:')
table(
    ['Statistic', 'Meaning'],
    [
        ['Votes Cast', 'Total votes received (digital + paper).'],
        ['Expected',   'Expected voters (attending + absentee from Voting Setup).'],
        ['Turnout %',  'Votes Cast ÷ Expected × 100. Shown green when ≥ 50%.'],
        ['Majority',   'Minimum votes needed to pass based on the configured threshold.'],
    ],
    col_widths=[4.0, 12.0]
)

warning('When the vote is marked complete, all voter phones show: "Thank you for your vote. The voting is now complete. The chairman will announce the results."')

h2('12b.  Voting Dashboard')
body(
    'The Voting Dashboard gives the chairman a read-only view of live congregational vote '
    'results. From the Voting Hub, click Voting Dashboard and enter the results password. '
    'It shows the vote question, live vote counts per answer, proportional bars, '
    'turnout %, majority threshold, and a green ✓ Majority badge when the threshold is reached.'
)

h2('12c.  Paper Ballot Entry — Congregational Vote')
body(
    'The same Paper Ballot Entry screen used for the election automatically switches to '
    'congregational vote mode when a congregational vote is open.'
)
numbered('The screen shows the vote question and answer options.')
numbered('Click the member\'s selected answer to highlight it.')
numbered('If the ballot is from an absentee voter, tick the Absentee ballot checkbox.')
numbered('Click Submit Paper Vote. The vote is recorded and the form resets.')


# ══════════════════════════════════════════════════════════════════
# 13. ROLES & RESPONSIBILITIES
# ══════════════════════════════════════════════════════════════════
h1('13.  Roles & Responsibilities')

table(
    ['Role', 'Responsibilities', 'Passwords Needed'],
    [
        ['Election Officer',       'Operates Election Setup, Round Control, and Voting Control; opens/closes voting; processes round transitions.', 'Admin, Election, Voting'],
        ['Chairman',               'Announces candidates and results; accesses Election Dashboard and Voting Dashboard; does not operate any control functions.', 'Results'],
        ['Paper Ballot Volunteer', 'Operates Paper Ballot Entry; collects and enters paper ballots for both election and congregational vote.', 'Paper Ballot'],
        ['Token Administrator',    'Generates and prints token cards before the meeting; distributes them at the door.', 'Tokens'],
        ['Door Steward',           'Distributes token cards to verified eligible voters only. No app password required.', '(none)'],
    ],
    col_widths=[3.8, 8.2, 4.0]
)

warning('Never share the results password with the election officer, or the admin/election password with the chairman. Keep roles separated.')


# ══════════════════════════════════════════════════════════════════
# 14. MEETING DAY CHECKLIST
# ══════════════════════════════════════════════════════════════════
h1('14.  Meeting Day Checklist')

h2('Before the Meeting — At Least One Week Prior')
for item in [
    'Log in to App Settings — change all six passwords from their defaults.',
    'Set congregation name, denomination, and upload logo if desired.',
    'Election Setup → Details: enter congregation name, meeting date, expected voters, absentee count.',
    'Election Setup → Elder: enter nominees, positions, votes per voter. Save.',
    'Election Setup → Deacon: enter nominees, positions, votes per voter (if Deacon election). Save.',
    'Voter Tokens: generate tokens (one per eligible voter). Print and cut token cards.',
    'Election Setup → Settings: confirm voter URL. Test it on a phone — confirm the token screen appears.',
    'Congregational Vote → Voting Setup: enter vote question, answer options, expected voters, majority threshold. Save.',
    'Print paper ballots for each office being held.',
    'Brief the chairman on the results password and how to access both dashboards.',
    'Brief the paper ballot volunteer on the paper ballot password and their screen.',
    'Store token cards securely until the meeting.',
]:
    p = doc.add_paragraph(); _space(p, 0, 4)
    p.add_run('☐   ' + item).font.size = Pt(11)

h2('Day of the Meeting')
for item in [
    'Confirm internet is working — test the voter URL on a phone.',
    'Log in to the admin hub on your device.',
    'Confirm all election setup is correct (nominees, voter count, voter URL).',
    'Open Election Dashboard on a separate device or browser tab for the chairman.',
    'Open Paper Ballot Entry on the volunteer\'s device.',
    'Distribute token cards at the entrance — one per eligible voter.',
    'Confirm paper ballot volunteer is logged in and the station is ready.',
]:
    p = doc.add_paragraph(); _space(p, 0, 4)
    p.add_run('☐   ' + item).font.size = Pt(11)

h2('After the Meeting')
for item in [
    'Chairman may screenshot Election Dashboard and Voting Dashboard for the minutes.',
    'Use Reset All Election Data (Election Setup → Settings → Danger Zone) to clear the election.',
    'Use Reset & New Vote (Voting Control) to clear the congregational vote data.',
    'Shred or store token cards securely.',
]:
    bullet(item)


# ══════════════════════════════════════════════════════════════════
# 15. SUGGESTED ORDER OF EVENTS
# ══════════════════════════════════════════════════════════════════
h1('15.  Suggested Order of Events')

body('Example: electing 2 Elders and 1 Deacon, then a congregational vote.')

h2('Elder Election')
for step in [
    'Chairman introduces the Elder nominees and explains the voting process.',
    'Chairman asks members to open the voter URL or scan the QR code on their token card.',
    'Election officer: Round Control → Start Round 1 Voting.',
    'Chairman: "Voting is now open. Please select your candidate(s) and submit."',
    'Paper ballot volunteer enters absentee ballots first (Absentee checkbox ticked), then remaining paper ballots.',
    'Election officer monitors Ballots In and Participation %. When satisfied: Close Voting → End Round → confirm.',
    'Election officer processes Round Transition: confirm elected candidates (Step 1); set advancing candidates if Round 2 is needed (Step 2); set votes per voter (Step 3).',
    'If all positions filled: Complete This Office. If not: Launch Next Round.',
    'Chairman announces results verbally (from Election Dashboard — vote counts not shown publicly).',
]:
    numbered(step)

h2('Deacon Election')
for step in [
    'Round Control now shows "Ready to Start" for Deacon. Chairman introduces nominees.',
    'Repeat the same flow as Elder (start round, open voting, monitor, close, transition, complete).',
    'Chairman announces elected Deacon(s).',
    'Voter phones show the thank-you message. Election officer sees the Election Complete screen.',
]:
    numbered(step)

h2('Congregational Vote')
for step in [
    'Chairman introduces the motion and reads it to the congregation.',
    'Voting Hub → Voting Control → Open Voting.',
    'Voter phones automatically switch to the congregational vote ballot (members who voted in the election do not need to re-enter their token).',
    'Paper ballot volunteer enters paper votes (screen switches to vote mode automatically).',
    'Election officer: Close Voting → Mark Complete (confirm).',
    'Chairman views Voting Dashboard with results password and announces the outcome.',
]:
    numbered(step)

tip('Always run a complete test election before the actual meeting. Use at least two phones to confirm the full flow from token entry to election complete.')


# ══════════════════════════════════════════════════════════════════
# 16. TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════
h1('16.  Troubleshooting')

table(
    ['Problem', 'Likely Cause', 'Solution'],
    [
        ['App password screen does not dismiss',
         'Wrong password entered.',
         'Re-enter the correct app password. If forgotten, contact your hosting administrator.'],
        ['Voter page shows "No Voting Configured"',
         'No election or vote is set up yet, or setup was saved with server offline.',
         'Complete Election Setup and save. If the voter URL is wrong, correct it in Settings.'],
        ['Voter cannot reach the voter page',
         'Wrong URL, or voter\'s phone has no internet connection.',
         'Confirm the URL on the token card matches your app URL. Confirm voter has internet access.'],
        ['"Failed to record vote" on voter\'s phone',
         'Phone lost internet connection momentarily.',
         'Have voter try again. If persistent, check that the app URL is reachable.'],
        ['"Invalid token" error',
         'Code mistyped or tokens were regenerated after printing.',
         'Double-check the 4-digit code. If tokens were regenerated, old cards are invalid — reprint.'],
        ['"Token already used" error',
         'Voter has already submitted a ballot this round, or a paper ballot was entered for their token.',
         'Voter has already voted this round. Verify with the paper ballot log if unsure.'],
        ['Voter phone stuck on "Voting Closed"',
         'Voting was closed in Round Control; voter was already on the ballot screen.',
         'Normal behaviour — tell voter to wait. Page updates automatically when next round opens.'],
        ['Vote counts not updating in Round Control',
         'Browser tab lost network connection.',
         'Refresh the browser tab. Counts update every 3 seconds when connected.'],
        ['Paper Ballot Entry form not appearing',
         'Voting is not currently open.',
         'Open voting in Round Control or Voting Control. Screen updates within 4 seconds.'],
        ['Congregational vote ballot not showing on phones',
         'Voting Control: voting is not open.',
         'Go to Voting Control → Open Voting. Phones update within 3 seconds.'],
        ['Login fails on mobile (password rejected)',
         'Mobile keyboard autocapitalised or added a space.',
         'Type the password carefully. All password fields disable autocorrect. Do not copy-paste.'],
        ['Forgotten password',
         'Password was changed and not recorded.',
         'Contact your hosting administrator to reset passwords.'],
    ],
    col_widths=[4.0, 4.5, 7.5]
)


# ══════════════════════════════════════════════════════════════════
# 17. QUICK REFERENCE
# ══════════════════════════════════════════════════════════════════
h1('17.  Quick Reference')

h2('Passwords — Fill In Before the Meeting')
table(
    ['Password', 'Default', 'Protects', 'Custom Password'],
    [
        ['Admin',        'election2024',   'App Settings',                      '____________________'],
        ['Election',     'election2024',   'Election Setup, Round Control',      '____________________'],
        ['Voting',       'voting2024',     'Voting Setup, Voting Control',       '____________________'],
        ['Results',      'results2024',    'Election Dashboard, Voting Dashboard','____________________'],
        ['Tokens',       'tokens2024',     'Voter Tokens',                       '____________________'],
        ['Paper Ballot', 'paperentry2024', 'Paper Ballot Entry',                 '____________________'],
    ],
    col_widths=[3.0, 3.0, 5.0, 5.0]
)

h2('App URL & Key Numbers')
table(
    ['', 'Value'],
    [
        ['Admin Hub URL',         '________________________________________________'],
        ['Voter Page URL',        '________________________________________________'],
        ['Expected Voters',       '________'],
        ['Absentee Votes',        '________'],
        ['Round 1 Majority',      'floor((Attending + Absentee) / 2) + 1  =  ________'],
        ['Round 2+ Majority',     'floor(Attending / 2) + 1  =  ________'],
        ['Congregational Majority','____________________________'],
    ],
    col_widths=[4.5, 11.5]
)

h2('Screen Navigation Summary')
table(
    ['Screen', 'Who', 'Password', 'Purpose'],
    [
        ['Home Screen',          'All officers',    'App password (session)', 'Navigation hub.'],
        ['App Settings',         'Elect. Officer',  'Admin',                  'Church identity and all passwords.'],
        ['Election Setup',       'Elect. Officer',  'Election',               'Configure nominees, voter count, URL.'],
        ['Round Control',        'Elect. Officer',  'Election (hub)',         'Open/close voting, monitor, transitions.'],
        ['Voter Tokens',         'Token Admin',     'Tokens',                 'Generate tokens, print cards.'],
        ['Voting Setup',         'Elect. Officer',  'Voting (hub)',           'Configure vote question and answers.'],
        ['Voting Control',       'Elect. Officer',  'Voting (hub)',           'Open/close congregational voting.'],
        ['Paper Ballot Entry',   'Volunteer',       'Paper Ballot',           'Enter paper votes — election or vote.'],
        ['Election Dashboard',   'Chairman',        'Results',                'Live election status and counts.'],
        ['Voting Dashboard',     'Chairman',        'Results',                'Live congregational vote results.'],
        ['Voter Page',           'All voters',      'Token code',            'Submit digital ballots.'],
    ],
    col_widths=[3.5, 3.0, 3.2, 6.3]
)

# Footer
fp = doc.add_paragraph(); _space(fp, 24, 0); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('Church Voting App — Election Officer Manual  |  votingapp.ca\nFor Reformed Church Congregational Meetings')
r.font.size = Pt(9); r.font.color.rgb = GREY

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'church_manual.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
