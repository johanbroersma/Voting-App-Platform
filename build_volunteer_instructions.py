#!/usr/bin/env python3
"""
Church Voting App — Volunteer (Paper Ballot Entry) Instructions
Generates volunteer_instructions.docx
Two half-page instruction cards per A4 page (cut along the dashed line).
Covers: paper ballot entry for office bearer election + congregational vote.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x27, 0x44)
GOLD   = RGBColor(0xb8, 0x94, 0x2a)
WHITE  = RGBColor(0xff, 0xff, 0xff)
GREY   = RGBColor(0x66, 0x66, 0x66)
LGREY  = RGBColor(0x44, 0x44, 0x44)
GREEN  = RGBColor(0x2d, 0x6a, 0x4f)
LGREEN = RGBColor(0xe8, 0xf5, 0xee)
AMBER  = RGBColor(0x92, 0x60, 0x10)
BROWN  = RGBColor(0x7c, 0x3d, 0x12)

doc = Document()

# ── A4 page, tight margins ────────────────────────────────────────────────────
for section in doc.sections:
    section.page_height = Inches(11.69)
    section.page_width  = Inches(8.27)
    section.top_margin    = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin   = Inches(0.45)
    section.right_margin  = Inches(0.45)

# ── XML helpers ───────────────────────────────────────────────────────────────
def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)

def set_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement('w:tcBorders')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}')
        if val:
            el.set(qn('w:val'),   val.get('val','single'))
            el.set(qn('w:sz'),    val.get('sz','4'))
            el.set(qn('w:space'),'0')
            el.set(qn('w:color'), val.get('color','auto'))
        else:
            el.set(qn('w:val'), 'none')
        tcB.append(el)
    tcPr.append(tcB)

def remove_table_borders(tbl):
    tblEl = tbl._tbl
    tblPr = tblEl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tblEl.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBdr.append(el)
    tblPr.append(tblBdr)

def set_row_height(row, height_inches, exact=True):
    tr    = row._tr
    trPr  = tr.get_or_add_trPr()
    trH   = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), str(int(height_inches * 1440)))
    trH.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    trPr.append(trH)

def para_space(p, before=0, after=0, line=14):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(line)

# ── Step row helper ───────────────────────────────────────────────────────────
def add_step(cell, number, heading, detail, circle_colour='1A2744', is_last=False):
    inner = cell.add_table(rows=1, cols=2)
    remove_table_borders(inner)
    inner.alignment = WD_TABLE_ALIGNMENT.LEFT
    inner.allow_autofit = False

    for row in inner.rows:
        row.cells[0].width = Inches(0.52)
        row.cells[1].width = Inches(6.85)

    circ = inner.rows[0].cells[0]
    circ.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade_cell(circ, circle_colour)
    cp = circ.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(cp, before=3, after=3)
    cr = cp.add_run(str(number))
    cr.bold = True
    cr.font.size = Pt(13)
    cr.font.color.rgb = WHITE

    text = inner.rows[0].cells[1]
    text.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    hp = text.paragraphs[0]
    para_space(hp, before=2, after=1)
    hr = hp.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(11)
    hr.font.color.rgb = NAVY

    dp = text.add_paragraph()
    para_space(dp, before=0, after=5 if not is_last else 2)
    dr = dp.add_run(detail)
    dr.font.size = Pt(9.5)
    dr.font.color.rgb = LGREY


# ── Divider bar helper ────────────────────────────────────────────────────────
def add_section_divider(cell, text, bg='E8F0FA', border='1A3A5C', text_colour=None):
    div = cell.add_table(rows=1, cols=1)
    remove_table_borders(div)
    dc = div.rows[0].cells[0]
    shade_cell(dc, bg)
    set_borders(dc,
        top   ={'val':'single','sz':'4','color':border},
        bottom={'val':'single','sz':'4','color':border},
        left  ={'val':'single','sz':'4','color':border},
        right ={'val':'single','sz':'4','color':border},
    )
    dp = dc.paragraphs[0]
    para_space(dp, before=2, after=2)
    dr = dp.add_run(text)
    dr.bold = True
    dr.font.size = Pt(9)
    if text_colour:
        dr.font.color.rgb = text_colour
    else:
        dr.font.color.rgb = NAVY


# ════════════════════════════════════════════════════════════════════════════════
# Build one instruction card
# ════════════════════════════════════════════════════════════════════════════════
def build_card(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    shade_cell(cell, 'FAFAF8')

    # Header bar
    hdr = cell.add_table(rows=1, cols=1)
    remove_table_borders(hdr)
    hc = hdr.rows[0].cells[0]
    shade_cell(hc, '1A2744')
    set_row_height(hdr.rows[0], 0.55)
    hp = hc.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(hp, before=5, after=3)
    ht = hp.add_run('PAPER BALLOT VOLUNTEER GUIDE')
    ht.bold = True
    ht.font.size = Pt(15)
    ht.font.color.rgb = WHITE
    hs = hc.add_paragraph()
    hs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(hs, before=0, after=5)
    hsr = hs.add_run('Entering paper ballots on behalf of members who cannot use the digital system')
    hsr.font.size = Pt(9)
    hsr.font.color.rgb = RGBColor(0xb8, 0xc8, 0xe0)

    # Gold divider
    gd = cell.add_table(rows=1, cols=1)
    remove_table_borders(gd)
    gc = gd.rows[0].cells[0]
    shade_cell(gc, 'B8942A')
    set_row_height(gd.rows[0], 0.05, exact=True)
    gc.paragraphs[0].paragraph_format.space_before = Pt(0)
    gc.paragraphs[0].paragraph_format.space_after  = Pt(0)

    # ── Getting started ──────────────────────────────────────────────────────
    add_section_divider(cell, 'GETTING STARTED — ACCESS PAPER BALLOT ENTRY', 'EEF2F8', '1A2744')

    add_step(cell, 1, 'Open the Admin Hub',
        'On the volunteer laptop, go to  http://localhost:8080/  (or the URL shown by the election officer). '
        'You will see the Church Voting App home screen.')

    add_step(cell, 2, 'Open Paper Ballot Entry',
        'Tap the  Paper Ballot Entry  tile on the home screen. '
        'Enter the paper ballot password when prompted (the officer will give you this). '
        'The station opens automatically to whichever vote is currently active.')

    # ── PART A: Office Bearer Election ───────────────────────────────────────
    add_section_divider(cell, 'PART A — OFFICE BEARER ELECTION  (Elder & Deacon)', 'EEF2F8', '1A2744')

    add_step(cell, 3, 'Check the Active Office',
        'The heading shows the current office (Elder or Deacon) and round. '
        'Only enter ballots for the office and round that is currently open.')

    add_step(cell, 4, 'Select Candidate(s)',
        'Tick the checkbox next to each candidate the member has selected on their paper ballot. '
        'Do not select more than the number of positions shown at the top of the form.')

    add_step(cell, 5, 'Mark Absentee if Applicable',
        'If the member is absent and submitted a written ballot, tick the  Absentee  checkbox '
        'before submitting. Absentee ballots are counted separately in Round 1 only.')

    add_step(cell, 6, 'Submit and Confirm',
        'Tap  Submit Paper Ballot. The ballot appears in the log below the form. '
        'Check the log entry matches the member\'s paper ballot, then destroy the paper ballot.',
        is_last=True)

    # ── PART B: Congregational Vote ──────────────────────────────────────────
    add_section_divider(cell, 'PART B — CONGREGATIONAL VOTE  (Motion / Resolution)', 'F0FAF0', '2D6A4F',
                        text_colour=GREEN)

    add_step(cell, 7, 'Screen Switches Automatically',
        'When the officer opens the congregational vote, the Paper Ballot Entry screen '
        'switches automatically. You do not need to navigate away or refresh.',
        circle_colour='2D6A4F')

    add_step(cell, 8, 'Select the Member\'s Answer',
        'The motion is shown at the top. Tap the answer button matching the member\'s choice '
        '("In favour" or "Not in favour", or the options configured). '
        'Tick  Absentee  if applicable.',
        circle_colour='2D6A4F')

    add_step(cell, 9, 'Submit and Log',
        'Tap  Submit Paper Ballot. The answer is recorded in the log. '
        'Each member may only have one ballot entered per vote — check the log before submitting.',
        circle_colour='2D6A4F', is_last=True)

    # Tips box
    tip_tbl = cell.add_table(rows=1, cols=1)
    remove_table_borders(tip_tbl)
    tip_cell = tip_tbl.rows[0].cells[0]
    shade_cell(tip_cell, 'FFF8E6')
    set_borders(tip_cell,
        top   ={'val':'single','sz':'6','color':'B8942A'},
        bottom={'val':'single','sz':'6','color':'B8942A'},
        left  ={'val':'single','sz':'6','color':'B8942A'},
        right ={'val':'single','sz':'6','color':'B8942A'},
    )
    tp = tip_cell.paragraphs[0]
    para_space(tp, before=3, after=2)
    tlbl = tp.add_run('IMPORTANT REMINDERS     ')
    tlbl.bold = True
    tlbl.font.size = Pt(9)
    tlbl.font.color.rgb = AMBER

    tips = [
        'Only enter a ballot when voting is open — the form will not accept entries when voting is closed.',
        'Do not submit more candidates than the positions available — the form will warn you.',
        'If you make a mistake, ask the election officer to delete the last entry from the ballot log.',
        'Keep the page open — the screen updates automatically when the officer moves to the next round or vote.',
        'Do not share the paper ballot password with voters.',
    ]
    for tip in tips:
        tp2 = tip_cell.add_paragraph()
        para_space(tp2, before=0, after=2)
        tr2 = tp2.add_run('▸  ' + tip)
        tr2.font.size = Pt(8.5)
        tr2.font.color.rgb = LGREY

    lp = tip_cell.add_paragraph()
    para_space(lp, before=0, after=3)

    # Footer
    fp = cell.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(fp, before=4, after=2)
    fr = fp.add_run('For the election officer only  ·  Keep this guide with the paper ballots')
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY
    fr.italic = True


# ════════════════════════════════════════════════════════════════════════════════
# Page layout — 2 cards per page with dashed cut line
# ════════════════════════════════════════════════════════════════════════════════
outer = doc.add_table(rows=3, cols=1)
remove_table_borders(outer)
outer.alignment = WD_TABLE_ALIGNMENT.LEFT
outer.allow_autofit = False

# Card 1
set_row_height(outer.rows[0], 5.3, exact=False)
build_card(outer.rows[0].cells[0])

# Cut line
set_row_height(outer.rows[1], 0.25, exact=True)
cut_cell = outer.rows[1].cells[0]
shade_cell(cut_cell, 'FFFFFF')
set_borders(cut_cell,
    top   ={'val':'dashed','sz':'6','color':'AAAAAA'},
    bottom={'val':'none'}, left={'val':'none'}, right={'val':'none'},
)
cp = cut_cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(cp, before=2, after=0)
cr = cp.add_run('✂   fold or cut here   ✂')
cr.font.size = Pt(8)
cr.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
cr.italic = True

# Card 2 (identical)
set_row_height(outer.rows[2], 5.3, exact=False)
build_card(outer.rows[2].cells[0])

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'volunteer_instructions.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
