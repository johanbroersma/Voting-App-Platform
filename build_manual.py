#!/usr/bin/env python3
"""
Church Voting App — Complete User Manual
Generates manual.docx   (run: python3 build_manual.py)
Covers v2.6.1: Elder/Deacon election + Congregational Vote, all 6 passwords.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x27, 0x44)
GOLD   = RGBColor(0xb8, 0x94, 0x2a)
ELDER  = RGBColor(0x7c, 0x3d, 0x12)
DEACON = RGBColor(0x1a, 0x3a, 0x5c)
GREEN  = RGBColor(0x2d, 0x6a, 0x4f)
RED    = RGBColor(0x9b, 0x23, 0x35)
GREY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xff, 0xff, 0xff)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=None, colour=None):
    run.bold   = bold
    run.italic = italic
    if size:   run.font.size = Pt(size)
    if colour: run.font.color.rgb = colour

def para_space(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def shade_cell(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val','single'))
            el.set(qn('w:sz'),    val.get('sz','4'))
            el.set(qn('w:space'),'0')
            el.set(qn('w:color'), val.get('color','auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def h1(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=18, after=6)
    p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '8')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '1a2744')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=12, after=4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=8, after=3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GREY
    return p

def body(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=0, after=5)
    p.add_run(text)
    p.runs[0].font.size = Pt(11)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    para_space(p, before=0, after=3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number')
    para_space(p, before=0, after=3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def callout(doc, label, text, bg_hex, border_hex):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, bg_hex)
    set_cell_border(cell,
        top    ={'val':'single','sz':'12','color':border_hex},
        bottom ={'val':'single','sz':'4', 'color':border_hex},
        left   ={'val':'single','sz':'12','color':border_hex},
        right  ={'val':'single','sz':'4', 'color':border_hex},
    )
    lp = cell.paragraphs[0]
    para_space(lp, before=2, after=2)
    lr = lp.add_run(f'{label}  ')
    lr.bold = True
    lr.font.size = Pt(10)
    tr = lp.add_run(text)
    tr.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def note(doc, text):    callout(doc, 'NOTE:',            text, 'E8F0FA', '1A3A5C')
def tip(doc, text):     callout(doc, 'TIP:',             text, 'F0FAF0', '2D6A4F')
def warning(doc, text): callout(doc, 'IMPORTANT:',       text, 'FFF4E0', 'B8942A')
def danger(doc, text):  callout(doc, 'ACTION REQUIRED:', text, 'FDF0F0', '9B2335')

def screenshot(doc, label, description):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    label_cell = tbl.cell(0, 0)
    shade_cell(label_cell, '1A2744')
    set_cell_border(label_cell,
        top={'val':'single','sz':'4','color':'1A2744'},
        left={'val':'single','sz':'4','color':'1A2744'},
        right={'val':'single','sz':'4','color':'1A2744'},
        bottom={'val':'single','sz':'4','color':'1A2744'},
    )
    lp = label_cell.paragraphs[0]
    para_space(lp, before=2, after=2)
    lr = lp.add_run(f'  SCREENSHOT — {label.upper()}')
    lr.bold = True
    lr.font.size = Pt(9)
    lr.font.color.rgb = WHITE
    desc_cell = tbl.cell(1, 0)
    shade_cell(desc_cell, 'F0F4FA')
    set_cell_border(desc_cell,
        top={'val':'single','sz':'4','color':'1A2744'},
        left={'val':'single','sz':'4','color':'1A2744'},
        right={'val':'single','sz':'4','color':'1A2744'},
        bottom={'val':'single','sz':'4','color':'1A2744'},
    )
    dp = desc_cell.paragraphs[0]
    para_space(dp, before=6, after=16)
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f'[ Insert screenshot here ]\n{description}')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(0x2a, 0x44, 0x6a)
    dr.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def data_table(doc, headers, rows, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=n)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, '1A2744')
        p = cell.paragraphs[0]
        para_space(p, before=2, after=2)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg = 'F8F6F1' if ri % 2 == 1 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            para_space(p, before=2, after=2)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
    if col_widths:
        for row2 in tbl.rows:
            for ci2, cell2 in enumerate(row2.cells):
                cell2.width = Cm(col_widths[ci2])
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


# ════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(cover, before=80, after=6)
cr = cover.add_run('Church Voting App')
cr.bold = True
cr.font.size = Pt(28)
cr.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(sub, before=6, after=10)
sr = sub.add_run('Complete User Manual')
sr.font.size = Pt(16)
sr.font.color.rgb = GOLD

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_space(ver, before=0, after=30)
vr = ver.add_run('Version 2.6.1')
vr.font.size = Pt(11)
vr.font.color.rgb = GREY

for line in [
    'Office Bearer Elections · Congregational Voting',
    'Reformed Church Congregational Meeting Software',
    '',
    'Confidential — For Election Officers & Chairman Use Only',
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_space(p, before=0, after=4)
    r = p.add_run(line)
    r.font.size = Pt(11)
    if 'Confidential' in line:
        r.bold = True
        r.font.color.rgb = NAVY
    else:
        r.font.color.rgb = GREY

pb = doc.add_paragraph()
pb.add_run().add_break(WD_BREAK.PAGE)


# ════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════
toc_title = doc.add_paragraph()
toc_title.paragraph_format.page_break_before = False
para_space(toc_title, before=0, after=8)
tr2 = toc_title.add_run('Table of Contents')
tr2.bold = True
tr2.font.size = Pt(16)
tr2.font.color.rgb = NAVY

toc_entries = [
    ('1.',     'System Overview'),
    ('2.',     'Technical Requirements & Setup'),
    ('   2a.', 'Running Locally'),
    ('   2b.', 'Cloud / Internet Hosting'),
    ('3.',     'Passwords & Security'),
    ('4.',     'First Run — Landing Page'),
    ('5.',     'Administration Setup (Election Setup)'),
    ('   5a.', 'Congregation Details'),
    ('   5b.', 'Elder Election Configuration'),
    ('   5c.', 'Deacon Election Configuration'),
    ('   5d.', 'Settings — URL, QR Code & Passwords'),
    ('6.',     'Voter Tokens'),
    ('7.',     'Congregational Vote Setup'),
    ('   7a.', 'Voting Hub'),
    ('   7b.', 'Voting Setup'),
    ('8.',     'During the Meeting — Round Control (Office Bearer Election)'),
    ('   8a.', 'Starting Round 1'),
    ('   8b.', 'Opening and Closing Voting'),
    ('   8c.', 'Monitoring Participation'),
    ('   8d.', 'Ending a Round & Round Transition'),
    ('   8e.', 'Completing an Office'),
    ('9.',     'Paper Ballot Entry — Office Bearer Election'),
    ('10.',    'Voter Page (vote.html)'),
    ('11.',    'Election Dashboard'),
    ('12.',    'Election Complete Screen'),
    ('13.',    'Congregational Vote — Running the Vote'),
    ('   13a.','Voting Control'),
    ('   13b.','Voting Dashboard'),
    ('   13c.','Paper Ballot Entry — Congregational Vote'),
    ('14.',    'Best Practices & Suggested Workflow'),
    ('15.',    'Complete Meeting Walkthrough'),
    ('16.',    'Troubleshooting'),
    ('17.',    'Quick Reference'),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    para_space(p, before=0, after=3)
    indent = len(num) - len(num.lstrip())
    p.paragraph_format.left_indent = Cm(indent * 0.2)
    r1 = p.add_run(num + '  ')
    r1.font.size = Pt(11)
    r1.bold = not num.startswith('   ')
    r1.font.color.rgb = NAVY if not num.startswith('   ') else GREY
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    r2.font.color.rgb = NAVY if not num.startswith('   ') else GREY

pb2 = doc.add_paragraph()
pb2.add_run().add_break(WD_BREAK.PAGE)


# ════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════════════════════
h1(doc, '1.  System Overview')

body(doc, (
    'The Church Voting App is a browser-based voting application for Reformed church '
    'congregational meetings. It supports office bearer elections (Elder and Deacon) '
    'and congregational votes (motions/resolutions). The system runs on a local Python '
    'server — no internet connection is required for basic use. All data is stored in '
    'a single shared file on the meeting laptop and accessed in real time by all '
    'connected devices.'
))

h2(doc, 'Key Files')
data_table(doc,
    ['File', 'Purpose', 'Who Uses It'],
    [
        ['server.py',           'Python HTTP server — serves the app and stores all data', 'Must stay running throughout the meeting'],
        ['index.html',          'Administration hub — all setup and control functions',    'Election Officer, Chairman'],
        ['vote.html',           'Voter ballot page — members cast their votes here',       'All eligible voters (phones/tablets)'],
        ['election_state.json', 'Shared data file — created automatically on first save',  '(Never opened directly)'],
    ],
    col_widths=[3.5, 8.0, 4.5]
)

h2(doc, 'Key Capabilities')
for feat in [
    'Multi-round office bearer elections with automatic majority detection',
    'Congregational vote (motion/resolution) with configurable majority threshold',
    'Digital voting via personal phones/tablets using unique 4-digit token codes',
    'Paper ballot entry station for members who prefer to vote on paper',
    'Absentee vote support for both election and congregational vote',
    'Six independent password levels protecting different areas of the system',
    'Real-time vote count updates — Round Control and Voting Control refresh every 3 seconds',
    'Voter page detects whether an election or congregational vote is active and shows the correct ballot automatically',
    'One token card per member covers all rounds, both offices, and the congregational vote',
    'No internet required — runs on a local Wi-Fi network. Can also be hosted publicly.',
]:
    bullet(doc, feat)

screenshot(doc, 'Landing Page — Home Screen',
    'Home screen showing the app password entry on first load. After login: four navigation '
    'tiles (Office Bearer Election, Congregational Vote, Voter Tokens, Paper Ballot Entry) '
    'with a warning banner if the default app password is still in use.')


# ════════════════════════════════════════════════════════════════════════
# 2. TECHNICAL REQUIREMENTS & SETUP
# ════════════════════════════════════════════════════════════════════════
h1(doc, '2.  Technical Requirements & Setup')

h2(doc, 'Administrator Device')
for item in [
    'Python 3.8 or later (pre-installed on most Macs; download from python.org for Windows)',
    'Any modern web browser (Chrome, Edge, Firefox, Safari — 2020 or later)',
    'server.py, index.html, and vote.html must be in the same folder',
    'Connected to a local Wi-Fi network (or internet for cloud hosting)',
]:
    bullet(doc, item)

h2(doc, 'Voter Devices')
for item in [
    'Any device with a modern web browser (phone, tablet, laptop)',
    'Connected to the same Wi-Fi network as the administrator\'s laptop',
    'No app installation required',
]:
    bullet(doc, item)

h2(doc, '2a.  Running Locally')
body(doc, 'Open a Terminal in the election folder and run:')
p = doc.add_paragraph()
para_space(p, before=2, after=6)
p.paragraph_format.left_indent = Cm(1.0)
r = p.add_run('python3 server.py')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = NAVY

body(doc, 'On startup the server prints:')
for line in [
    'Admin / local:   http://localhost:8080/',
    'Voter devices:   http://192.168.x.x:8080/vote.html',
]:
    p = doc.add_paragraph()
    para_space(p, before=0, after=2)
    p.paragraph_format.left_indent = Cm(1.0)
    p.add_run(line).font.size = Pt(10)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

numbered(doc, 'Open http://localhost:8080/ in the admin browser on the laptop.')
numbered(doc, 'Note the voter URL printed on startup — you will enter it in Settings.')
numbered(doc, 'Voters open http://192.168.x.x:8080/vote.html on their phones.')

warning(doc, 'The server must stay running throughout the meeting. Do not close the Terminal window.')

h2(doc, '2b.  Cloud / Internet Hosting (Render.com)')
body(doc, (
    'For remote or large-congregation use, the app can be deployed publicly. '
    'Voters access the ballot page over the internet without needing to join the meeting Wi-Fi.'
))
for step in [
    'Push the project folder to a GitHub repository.',
    'Go to render.com → New → Blueprint. Connect the GitHub repository.',
    'Render reads render.yaml automatically and provisions the service.',
    'Your public URL appears in the Render dashboard (e.g. https://church-vote.onrender.com).',
    'In Election Setup → Settings, save the public /vote.html address as the voter URL.',
]:
    numbered(doc, step)

note(doc, (
    'A persistent disk (Render Starter plan, ~$7/month) keeps election_state.json if the '
    'service restarts. The free plan is acceptable for a single meeting day.'
))


# ════════════════════════════════════════════════════════════════════════
# 3. PASSWORDS & SECURITY
# ════════════════════════════════════════════════════════════════════════
h1(doc, '3.  Passwords & Security')

body(doc, (
    'The system uses six independent passwords. Each password protects a specific area '
    'so that the election officer, chairman, paper ballot volunteer, and token administrator '
    'only have access to their respective functions. Passwords are stored as SHA-256 '
    'cryptographic hashes — the plaintext password is never saved anywhere.'
))

data_table(doc,
    ['Password', 'Protects', 'Default', 'Held By'],
    [
        ['App (Landing) Password',  'Entire admin hub — required on first load each session',    'votevote2024',  'All officers'],
        ['Admin Password',          'Election Setup and Round Control',                           'election2024',  'Election Officer only'],
        ['Voting Password',         'Voting Setup and Voting Control (Congregational Vote)',      'voting2024',    'Election Officer only'],
        ['Results Password',        'Election Dashboard (confidential vote counts)',              'results2024',   'Chairman only'],
        ['Tokens Password',         'Voter Tokens section (token generation and printing)',       'tokens2024',    'Token administrator'],
        ['Paper Ballot Password',   'Paper Ballot Entry station (both election and vote)',        'paperentry2024','Paper ballot volunteer'],
    ],
    col_widths=[4.0, 5.5, 3.0, 3.5]
)

danger(doc, (
    'Change ALL passwords from their defaults before the meeting. The defaults above are '
    'publicly documented. Go to Election Setup → Settings, Voting Setup, and the Voter '
    'Tokens section to change each one.'
))

body(doc, (
    'The landing page password gates access to the entire admin hub (index.html) for the '
    'current browser session. Once entered, the session stays authenticated until the browser '
    'is closed. A yellow warning banner appears on the home screen if the default landing '
    'password is still in use.'
))

warning(doc, (
    'Keep passwords separate. The chairman does not need the admin password. The paper ballot '
    'volunteer does not need the results password or admin password. Distribute each password '
    'only to the person who needs it.'
))

note(doc, (
    'All password fields have autocapitalize, autocorrect, and spellcheck disabled to prevent '
    'mobile keyboard interference. Type passwords carefully and avoid copy-pasting from '
    'other apps, which may introduce invisible characters.'
))


# ════════════════════════════════════════════════════════════════════════
# 4. FIRST RUN — LANDING PAGE
# ════════════════════════════════════════════════════════════════════════
h1(doc, '4.  First Run — Landing Page')

body(doc, (
    'When index.html is first opened in the browser, the system shows an App Password screen. '
    'Enter the landing password (default: votevote2024) and click Log In. This must be '
    'done once per browser session — the session remains authenticated until the browser is closed.'
))

screenshot(doc, 'Landing Page — App Password Screen',
    'Full-screen password entry with cross/church logo at the top. "Church Voting App" heading. '
    '"Enter App Password" label, password input field, and "Log In" button. '
    'Footer text: "Authorised personnel only."')

body(doc, 'After logging in, the home screen shows four navigation tiles:')
data_table(doc,
    ['Tile', 'Destination', 'Password Required'],
    [
        ['Office Bearer Election', 'Election Hub — Setup, Round Control, Dashboard',    'Admin password (at hub)'],
        ['Congregational Vote',    'Voting Hub — Voting Setup, Voting Control, Dashboard','Voting password (at hub)'],
        ['Voter Tokens',           'Token generation, printing, and QR code',           'Tokens password'],
        ['Paper Ballot Entry',     'Paper ballot entry station',                         'Paper ballot password'],
    ],
    col_widths=[4.5, 6.5, 5.0]
)

body(doc, (
    'If the app password is still the default (votevote2024), a yellow warning banner '
    'appears at the top of the home screen reminding officers to change it. '
    'The banner disappears as soon as the password is changed in Election Setup → Settings.'
))

screenshot(doc, 'Home Screen',
    'Four large navigation tiles in a 2x2 grid. Top-left: "Office Bearer Election" (navy). '
    'Top-right: "Congregational Vote" (navy). Bottom-left: "Voter Tokens" (navy). '
    'Bottom-right: "Paper Ballot Entry" (navy). Yellow warning banner at the top if the '
    'default app password is still in use.')


# ════════════════════════════════════════════════════════════════════════
# 5. ADMINISTRATION SETUP — ELECTION SETUP
# ════════════════════════════════════════════════════════════════════════
h1(doc, '5.  Administration Setup — Election Setup')

body(doc, (
    'Election Setup is where the election officer configures the office bearer election '
    'before the meeting. From the home screen, click Office Bearer Election, then click '
    'Election Setup. Enter the admin password to access.'
))

body(doc, 'Election Setup is organised into four tabs: Details, Elder Election, Deacon Election, and Settings.')

note(doc, (
    'Configure only the offices being voted on this meeting. '
    'If only Elders are being elected, leave the Deacon tab empty — the system skips unconfigured offices.'
))

screenshot(doc, 'Election Setup — Tabs',
    'Four tab buttons at the top: Details, Elder Election, Deacon Election, Settings. '
    'Each tab is highlighted when active. A "▶ Round Control" shortcut button appears '
    'in the header once at least one office is configured.')

h2(doc, '5a.  Congregation Details')
body(doc, 'Click the Details tab and fill in:')
data_table(doc,
    ['Field', 'Description', 'Example'],
    [
        ['Congregation Name',           'Full name shown on token cards, voter page header, and dashboard.',  'First Reformed Church'],
        ['Meeting Date',                'Date of the meeting. Shown on voter page header.',                   '2026-04-22'],
        ['Expected Voters (attending)', 'Eligible male voters attending in person. Used for majority and participation %.', '80'],
        ['Absentee Votes (Round 1)',    'Pre-meeting paper ballots received. Counted only in Round 1.',       '5'],
    ],
    col_widths=[4.5, 8.0, 3.5]
)
body(doc, (
    'Below these fields the system displays the Majority Required threshold(s). '
    'Majority = floor(n/2) + 1. Round 1 uses Expected Voters + Absentee Votes. '
    'Round 2 and later use Expected Voters only.'
))
tip(doc, 'Fields save when you move to the next field. Click Save Details to save manually.')
warning(doc, 'Enter absentee count before the meeting. Each absentee ballot is still entered individually during the Round 1 voting period via Paper Ballot Entry with the Absentee checkbox ticked.')

h2(doc, '5b.  Elder Election Configuration')
body(doc, 'Click the Elder Election tab and enter:')
data_table(doc,
    ['Field', 'Description'],
    [
        ['Nominees',             'One name per line. All men nominated for the office of Elder.'],
        ['Positions to Fill',    'Number of Elder positions to be filled in this election.'],
        ['Votes per Voter (R1)', 'How many candidates each voter may select in Round 1. Typically equal to positions.'],
    ],
    col_widths=[4.5, 11.5]
)
body(doc, 'Click Save Elder Setup when done.')
warning(doc, 'Saving setup resets all in-progress election data for that office. Never click Save during an active round.')

h2(doc, '5c.  Deacon Election Configuration')
body(doc, (
    'Click the Deacon Election tab and enter nominees, positions, and votes per voter — '
    'same process as Elder. Click Save Deacon Setup. '
    'Leave blank if no Deacon election is being held.'
))

h3(doc, 'Printing Paper Ballots')
bullet(doc, 'Set the number of Elder ballots and click Print Elder Ballots.')
bullet(doc, 'Set the number of Deacon ballots and click Print Deacon Ballots.')
bullet(doc, 'Each ballot shows: office name, congregation, positions to fill, and all nominees with checkboxes.')
note(doc, 'Paper ballots do not show round information — the same printed ballots can be used in all rounds. The chairman verbally announces which candidates are on the ballot each round.')

h2(doc, '5d.  Settings — URL, QR Code & Passwords')
h3(doc, 'Voting Page URL')
body(doc, (
    'Enter the full URL that voters will use to access vote.html. '
    'For local network: http://192.168.x.x:8080/vote.html. '
    'Click Save URL. The QR code updates automatically.'
))
warning(doc, 'Save the voter URL BEFORE printing token cards. The QR code and URL printed on each card are generated from this saved value.')

h3(doc, 'Change App (Landing) Password')
body(doc, 'Enter current password, new password, and confirm. Click Update App Password. This affects the landing page login for all officers.')

h3(doc, 'Change Admin Password')
body(doc, 'Enter current admin password, new password, and confirm. Click Update Admin Password.')

h3(doc, 'Change Results Password')
body(doc, 'Enter current results password, new password, and confirm. Click Update Results Password. Share this password only with the chairman.')

h3(doc, 'Change Paper Ballot Password')
body(doc, 'Allows updating the paper ballot station password from within Election Setup Settings (as an alternative to Voting Setup).')

h3(doc, 'Reset All Election Data')
body(doc, (
    'The red Reset All Election Data button (Danger Zone section) permanently erases all '
    'nominees, ballots, tokens, results, and resets passwords to defaults. '
    'Two confirmation prompts prevent accidental use.'
))
danger(doc, 'Reset All Election Data cannot be undone. Use only to start fresh before a new election.')

screenshot(doc, 'Election Setup — Settings Tab',
    'Sections: Voting Page URL (current URL display, input field, Save and Clear buttons), '
    'QR Code (generated QR image and Print QR button), Change App Password card, '
    'Change Admin Password card, Change Results Password card, '
    'Change Paper Ballot Password card, Danger Zone (red Reset button).')


# ════════════════════════════════════════════════════════════════════════
# 6. VOTER TOKENS
# ════════════════════════════════════════════════════════════════════════
h1(doc, '6.  Voter Tokens')

body(doc, (
    'The Voter Tokens section is accessed from the home screen. Click Voter Tokens and '
    'enter the tokens password (default: tokens2024). '
    'Each eligible voter receives one token card. '
    'A single token covers all election rounds, both offices (Elder and Deacon), and the '
    'congregational vote.'
))

h2(doc, 'Generating Tokens')
for step in [
    'Enter the number of tokens (equal to the total number of eligible voters, e.g. 85).',
    'Click Generate Tokens. The system creates unique 4-digit codes.',
    'The token grid displays all codes. Used tokens are shown with a grey badge.',
]:
    numbered(doc, step)

warning(doc, (
    'Generating new tokens replaces all existing ones. Do this only once before the meeting. '
    'Regenerating tokens during a live election invalidates all already-distributed cards.'
))

h2(doc, 'Printing Token Cards')
body(doc, (
    'Click Print Token Cards. Each card shows: congregation name, the 4-digit code, a QR code '
    'linking to vote.html, and the voter URL. Print on card stock and cut into individual cards. '
    'Distribute one card per eligible voter at the entrance.'
))

screenshot(doc, 'Token Cards — Print Preview',
    '4-column grid of token cards on a cream background. Each card shows: congregation name '
    'at top, large 4-digit code in navy, a QR code image, and the voter URL at the bottom. '
    'A "fold or cut" dashed line separates rows.')

h2(doc, 'Voting Page URL & QR Code')
body(doc, (
    'This section also shows the current voter URL and QR code. '
    'Members can scan the QR code on any screen to open the voter page without typing the URL.'
))

h2(doc, 'Change Tokens Password')
body(doc, 'Enter current tokens password, new password, and confirm. Click Update Tokens Password.')


# ════════════════════════════════════════════════════════════════════════
# 7. CONGREGATIONAL VOTE SETUP
# ════════════════════════════════════════════════════════════════════════
h1(doc, '7.  Congregational Vote Setup')

body(doc, (
    'The Congregational Vote feature handles motions and resolutions that require a vote '
    'by the congregation. It is completely independent of the office bearer election — '
    'it can be run before, after, or without an election. '
    'From the home screen, click Congregational Vote and enter the voting password.'
))

h2(doc, '7a.  Voting Hub')
body(doc, 'The Voting Hub is the navigation centre for the congregational vote. It shows:')
bullet(doc, 'Status overview card (configured/open/closed/complete)')
bullet(doc, 'Navigation to Voting Setup, Voting Control, and Voting Dashboard')
bullet(doc, 'The voter URL for the congregational vote (same URL as the office bearer election)')

screenshot(doc, 'Voting Hub',
    'Dark navy header "Congregational Vote" with church name. Status overview card below. '
    'Three navigation tiles: Voting Setup (gear icon), Voting Control (controls icon), '
    'Voting Dashboard (chart icon). Voter URL displayed at the bottom.')

h2(doc, '7b.  Voting Setup')
body(doc, (
    'Click Voting Setup in the Voting Hub to configure the vote. '
    'The voting password is required.'
))
data_table(doc,
    ['Field', 'Description', 'Example'],
    [
        ['Vote Question',       'The motion or resolution text that members will vote on.',           'Do you agree with the proposed resolution regarding the building fund?'],
        ['Answer Options',      'The choices members select from. At least two are required. Default: "In favour" / "Not in favour". Add or remove options as needed.', 'In favour\nNot in favour'],
        ['Expected Voters',     'Eligible voters for this vote. Defaults to the election value if left at 0.', '80'],
        ['Absentee Voters',     'Number of absentee ballots included in the majority calculation.',  '5'],
        ['Majority Threshold',  '50%+1 (simple majority), 75% of eligible voters, or a custom percentage.', '50%+1'],
    ],
    col_widths=[4.0, 8.5, 3.5]
)
body(doc, 'Click Save Voting Setup when done. A green confirmation appears.')
tip(doc, 'The "▶ Voting Control" shortcut button appears in the header once the vote is configured.')

warning(doc, (
    'Saving Voting Setup resets all votes and ballot data for the current vote. '
    'Never save setup while voting is open or after votes have been cast.'
))

screenshot(doc, 'Voting Setup',
    'Header "Voting Setup" with back button. Cards: Vote Question (textarea), '
    'Answer Options (list with add/remove buttons), Expected Voters (number input), '
    'Absentee Voters (number input), Majority Threshold (radio buttons: 50%+1 / 75% / Custom), '
    'Change Paper Ballot Password card. Save Voting Setup button at the bottom.')


# ════════════════════════════════════════════════════════════════════════
# 8. ROUND CONTROL — OFFICE BEARER ELECTION
# ════════════════════════════════════════════════════════════════════════
h1(doc, '8.  During the Meeting — Round Control (Office Bearer Election)')

body(doc, (
    'Round Control is the central panel used during the office bearer election. '
    'From the Office Bearer Election hub, click Round Control. '
    'The admin password is required.'
))

screenshot(doc, 'Round Control — Active Round',
    'Coloured header (warm brown for Elder, deep blue for Deacon) with office name and round '
    'number as subtitle. Sequence bar showing election progress (Elder → Deacon → Complete). '
    'Four stat boxes: Ballots In, Paper Ballots, Expected Voters, Participation %. '
    'Control bar with green "▶ Open Voting", red "■ Close Voting", and navy "📊 End Round" buttons '
    'plus a voting status indicator. Live Results list below with vote bars.')

h2(doc, '8a.  Starting Round 1')
body(doc, (
    'When no election is in progress, Round Control shows a "Ready to Start" panel for the '
    'first configured office. It displays the number of nominees, positions, and votes per '
    'voter. Click Start Round 1 Voting.'
))
tip(doc, 'Do not click Start Round 1 until the chairman has introduced the candidates and invited members to vote. Once clicked, the ballot is immediately available on voters\' phones.')

h2(doc, '8b.  Opening and Closing Voting')
data_table(doc,
    ['Button', 'Action', 'When to Use'],
    [
        ['▶ Open Voting',  'Opens voting — members can now submit ballots',              'At the start of each voting period'],
        ['■ Close Voting', 'Closes voting — no further ballots accepted',                'When sufficient participation is reached'],
        ['📊 End Round',   'Triggers confirmation dialog then opens Round Transition',   'After voting is closed'],
    ],
    col_widths=[3.5, 6.5, 6.0]
)
warning(doc, 'End Round requires voting to be closed first. A confirmation dialog shows the ballot count — click Confirm to proceed.')
body(doc, 'When voting is closed, all voter screens update within 3 seconds to show "Voting Round Closed — please wait for further instructions."')

h2(doc, '8c.  Monitoring Participation')
body(doc, 'The stat boxes update every 3 seconds automatically:')
data_table(doc,
    ['Statistic', 'Meaning'],
    [
        ['Ballots In',       'Total ballots this round (digital + paper combined)'],
        ['Paper Ballots',    'Ballots entered via the Paper Ballot Entry station'],
        ['Absentee',         'Paper ballots flagged as absentee (Round 1 only — hidden in Round 2+)'],
        ['Expected Voters',  'Attending + absentee in Round 1; attending only in Round 2+'],
        ['Participation %',  'Ballots In ÷ Expected Voters × 100. Shown green when ≥ 50%.'],
    ],
    col_widths=[4.5, 11.5]
)
note(doc, 'Live Results (vote counts per candidate) are visible only on Round Control and the Election Dashboard. They are not shown to voters or on any public screen.')

h2(doc, '8d.  Ending a Round & Round Transition')
body(doc, 'After clicking End Round and confirming, the Round Transition screen opens with four sections:')
bullet(doc, 'Final Results — all candidates ranked by votes received this round')
bullet(doc, 'Step 1 — Mark as Elected: candidates who reached majority are pre-checked')
bullet(doc, 'Step 2 — Select Who Advances: choose which non-elected candidates continue to the next round')
bullet(doc, 'Step 3 — Next Round Settings: set votes per voter for the next round')

body(doc, 'Step-by-step process:')
for step in [
    'Review Final Results. Candidates meeting the majority threshold are pre-checked in Step 1 with a gold "majority" badge.',
    'Adjust Step 1 checkboxes if needed. Uncheck to override an auto-suggestion; check to manually elect a candidate.',
    'If positions remain unfilled: Step 2 shows non-elected candidates (all advancing by default). Uncheck to exclude a candidate from the next round.',
    'Step 3 auto-fills votes per voter as the number of remaining positions. Adjust if needed.',
    'If all positions are filled: a green banner appears. Steps 2 and 3 are hidden. Only Complete This Office is available.',
    'Click Launch Next Round (opens Round Control with voting closed, so chairman can announce the round before opening) OR click Complete This Office.',
]:
    numbered(doc, step)

screenshot(doc, 'Round Transition Screen',
    'Final Results list with vote bars. Step 1 checkboxes — majority candidates pre-checked '
    'with gold badge. Step 2 checkboxes for advancing candidates. Step 3 votes-per-voter input. '
    'Green banner when all positions filled. "Launch Next Round" (green) and '
    '"Complete This Office" (navy) buttons at the bottom.')

h2(doc, '8e.  Completing an Office')
body(doc, 'When all positions are filled, click Complete This Office. The system:')
for step in [
    'Records elected candidates for this office.',
    'Marks the office as complete.',
    'If the other office is configured and not started: shows Round Control in "Ready to Start" mode for that office.',
    'If both offices are complete: shows the Election Complete screen.',
]:
    numbered(doc, step)


# ════════════════════════════════════════════════════════════════════════
# 9. PAPER BALLOT ENTRY — OFFICE BEARER ELECTION
# ════════════════════════════════════════════════════════════════════════
h1(doc, '9.  Paper Ballot Entry — Office Bearer Election')

body(doc, (
    'The Paper Ballot Entry screen is used by a dedicated volunteer to enter votes for '
    'members who voted on paper. It is accessed from the home screen — the paper ballot '
    'password is required (default: paperentry2024). '
    'The screen automatically detects whether an election or a congregational vote is '
    'active and shows the appropriate entry form.'
))

warning(doc, 'Open Paper Ballot Entry on a separate device or a separate browser tab. This screen is for a dedicated volunteer — not the election officer.')

h2(doc, 'Dynamic Behaviour')
data_table(doc,
    ['Condition', 'What the Screen Shows'],
    [
        ['No voting open',                   '"Voting is currently closed" message with current status'],
        ['Office bearer election open',      'Candidate list with checkbox rows and Submit Paper Ballot button'],
        ['Congregational vote open',         'Vote question with answer options and Submit Paper Vote button'],
        ['Round or vote changes',            'Screen auto-refreshes within 4 seconds and shows the new form'],
    ],
    col_widths=[5.5, 10.5]
)

h2(doc, 'Entering an Election Paper Ballot')
for step in [
    'Verify that the correct office and round are shown at the top.',
    'Click each candidate\'s name that was marked on the paper ballot. Selected candidates are highlighted with a checkmark.',
    '(Round 1 only) If the ballot is from an absentee voter, tick the "Absentee ballot" checkbox before submitting.',
    'Click Submit Paper Ballot. The vote is recorded and the form resets.',
    'The submitted ballot appears in the log at the bottom.',
]:
    numbered(doc, step)

h2(doc, 'Editing or Deleting a Ballot')
bullet(doc, 'To edit: click the ✏️ Edit button next to a ballot in the log. Make corrections and click Update Paper Ballot.')
bullet(doc, 'To delete the most recent entry: click Delete Last Entry (red button) and confirm.')
note(doc, 'Only the most recently submitted ballot can be deleted. Use the Edit button to correct earlier ballots.')

screenshot(doc, 'Paper Ballot Entry — Election',
    'Office badge at top showing current office and round. Selectable candidate rows. '
    'Absentee ballot checkbox (visible in Round 1 only). Green Submit Paper Ballot button. '
    'Red Delete Last Entry button below. Log of submitted ballots at the bottom, with '
    'edit buttons and Absentee badges on absentee entries.')


# ════════════════════════════════════════════════════════════════════════
# 10. VOTER PAGE (vote.html)
# ════════════════════════════════════════════════════════════════════════
h1(doc, '10.  Voter Page (vote.html)')

body(doc, (
    'The voter page is the single URL that all voters use throughout the meeting — '
    'for both the office bearer election and the congregational vote. '
    'The page automatically detects which type of voting is currently active and '
    'presents the appropriate ballot. No manual navigation or page refresh is needed.'
))

h2(doc, 'Accessing the Page')
bullet(doc, 'QR Code: Scan the QR code on the token card with the phone camera.')
bullet(doc, 'URL: Type the voter URL from the token card into the browser.')
bullet(doc, 'The page header shows the congregation name, meeting date, and a live status pill.')

h2(doc, 'Office Bearer Election — Voting Flow')
for step in [
    'The page shows a numeric token input. Enter the 4-digit code from the token card and tap Submit Token.',
    'The ballot appears showing the office (Elder or Deacon), round number, candidates, and how many selections are permitted.',
    'Tap each candidate to select. A coloured highlight and checkmark confirm the choice.',
    'Tap Submit My Vote. A "Vote Submitted!" confirmation appears with a summary of who was voted for.',
    'Keep the page open. If another round or office opens, a notification appears — tap Vote Now to continue.',
    'The same token card is used for all rounds and both offices. No re-entry is needed between rounds for the same office.',
]:
    numbered(doc, step)

screenshot(doc, 'Voter Page — Token Entry',
    'White card on navy background. Header: congregation name, meeting date, "Voting Open" '
    'pill. "Enter Your Voting Token" heading. Large 4-digit numeric input field '
    '(telephone keypad on mobile). Submit Token button.')

screenshot(doc, 'Voter Page — Election Ballot',
    '"Office of Elder — Round 1" badge. "Select up to 2 candidates" instruction. '
    'Pip indicators showing 1 of 2 selected. Candidate rows with checkmarks on selected ones. '
    'Submit My Vote button. "Your vote is anonymous" note.')

screenshot(doc, 'Voter Page — Vote Submitted',
    '"Vote Submitted!" with large green checkmark. Collapsible "Who you voted for" summary. '
    'Notification card below: "Round 2 voting is now open — Vote Now →" button (appears '
    'when the next round opens).')

h2(doc, 'Congregational Vote — Voting Flow')
body(doc, (
    'After the office bearer election concludes, or independently if only a congregational '
    'vote is held, the voter page automatically switches to congregational vote mode.'
))
for step in [
    'If the voter already entered their token for the election, the page transitions automatically to the congregational vote without requiring re-entry.',
    'If no election was held, the voter enters their 4-digit token as usual.',
    'The page shows a "Congregational Vote" badge, the vote question, and the answer options.',
    'Tap the desired answer. Tap Cast My Vote.',
    'A "Vote Recorded" confirmation appears showing the question and the selected answer.',
]:
    numbered(doc, step)

screenshot(doc, 'Voter Page — Congregational Vote Ballot',
    '"Congregational Vote" navy badge. Vote question displayed in a bordered card. '
    '"Select your answer:" label. Answer options as large tappable radio rows. '
    '"Cast My Vote" button (disabled until an answer is selected).')

screenshot(doc, 'Voter Page — Congregational Vote Confirmed',
    '"Vote Recorded" heading with green checkmark. "Thank you — your vote has been recorded." '
    'Question card and a green-bordered "Your vote" card showing the selected answer.')

h2(doc, 'All Voter Screen States')
data_table(doc,
    ['State', 'When It Appears'],
    [
        ['Token Entry',            'Voting is open and voter has not yet entered their token this session'],
        ['Election Ballot',        'Token entered; office bearer election is open'],
        ['Congregational Ballot',  'Token entered; congregational vote is open'],
        ['Vote Submitted (elect.)','"Vote Submitted!" — election ballot confirmed. Polls for next round.'],
        ['Vote Recorded (vote)',   '"Vote Recorded" — congregational vote confirmed'],
        ['Voting Closed',          'Round or vote closed while voter is on token/ballot screen'],
        ['Waiting',                'Voting not yet open or between rounds'],
        ['Election Complete',      'All election offices done — thank-you message shown'],
        ['Vote Complete',          'Congregational vote marked complete — thank-you message shown'],
        ['No Voting Configured',   'No election or vote is set up, or server is not reachable'],
    ],
    col_widths=[5.0, 11.0]
)

note(doc, 'One token covers everything: all election rounds, both offices, and the congregational vote. Members never need a new token card during the meeting.')
warning(doc, 'Votes are final once submitted. Digital ballots cannot be changed. Only paper ballots can be edited via Paper Ballot Entry.')


# ════════════════════════════════════════════════════════════════════════
# 11. ELECTION DASHBOARD
# ════════════════════════════════════════════════════════════════════════
h1(doc, '11.  Election Dashboard')

body(doc, (
    'The Election Dashboard is a password-protected, read-only view for the chairman. '
    'It shows live election status and confidential vote counts, refreshing every 3 seconds. '
    'From the Office Bearer Election hub, click Election Dashboard and enter the results password.'
))

screenshot(doc, 'Election Dashboard',
    'Navy header "Election Dashboard" with congregation name and meeting date. '
    'Status card showing current office, round, voting status pill (green/amber), '
    'majority required, votes per voter, and candidate name chips. '
    'Info strip: Expected Voters, Absentee Votes pills. '
    'Results section below showing all candidates with vote bars, vote counts, '
    'and ★ majority / ✓ Elected badges.')

h2(doc, 'Status Section')
data_table(doc,
    ['Information', 'Description'],
    [
        ['Office & Round',        'Which office and round is currently active'],
        ['Status Pill',           'Voting Open (green) / Voting Closed (amber) / Configuring (blue)'],
        ['Majority Required',     'Minimum votes for election this round (round-aware)'],
        ['Votes per Voter',       'How many candidates each voter may select this round'],
        ['Candidates This Round', 'Names of all candidates on the current ballot'],
    ],
    col_widths=[4.5, 11.5]
)

h2(doc, 'Results Section')
body(doc, 'The results section adapts to the election state:')
bullet(doc, 'Before Round 1: shows nominees, positions to fill, and Round 1 majority')
bullet(doc, 'During a round (voting open): shows live vote counts and bars with "● Voting Open" label')
bullet(doc, 'After End Round, before transition: shows a "Round N Results — Awaiting transition" card with full counts')
bullet(doc, 'After transition: shows committed historical results for each completed round')

tip(doc, 'The "Awaiting transition" card lets the chairman review results and prepare their announcement before the election officer processes the round transition.')

warning(doc, 'The dashboard shows actual vote counts. Keep the results password confidential and never leave the dashboard visible on an unattended shared screen.')


# ════════════════════════════════════════════════════════════════════════
# 12. ELECTION COMPLETE SCREEN
# ════════════════════════════════════════════════════════════════════════
h1(doc, '12.  Election Complete Screen')

body(doc, (
    'After the election officer clicks Complete This Office for the last configured office, '
    'the system shows the Election Complete screen. This screen is for the election officer only.'
))
body(doc, 'The screen displays:')
bullet(doc, 'A navy header: "Election Complete" with the meeting date')
bullet(doc, 'A card for each configured office listing the elected candidates by name only (no vote counts)')
bullet(doc, 'A "Back to Home" button to return to the landing screen')
note(doc, 'Vote counts are intentionally not shown here. The chairman announces results verbally and may view counts on the Election Dashboard.')
body(doc, 'All voters\' phones simultaneously show: "Thank you for participating in the [Church Name] office bearer election. The election has been completed. The chairman will announce the results."')

screenshot(doc, 'Election Complete Screen',
    'Navy header "Election Complete" with meeting date. "Office of Elder" card listing '
    'elected names with coloured dot indicators. "Office of Deacon" card below. '
    '"Back to Home" button at the bottom.')


# ════════════════════════════════════════════════════════════════════════
# 13. CONGREGATIONAL VOTE — RUNNING THE VOTE
# ════════════════════════════════════════════════════════════════════════
h1(doc, '13.  Congregational Vote — Running the Vote')

body(doc, (
    'Once the vote is configured in Voting Setup, the election officer manages it through '
    'Voting Control. The chairman monitors results on the Voting Dashboard. '
    'Both are accessed from the Voting Hub (Congregational Vote → enter voting password).'
))

h2(doc, '13a.  Voting Control')
body(doc, (
    'Voting Control is the officer\'s panel for the congregational vote. '
    'It mirrors the Round Control layout for consistency.'
))

screenshot(doc, 'Voting Control',
    'Coloured header (navy when closed, green when open, gold when complete) "Voting Control". '
    'Subtitle: "[Church] · Congregational Vote". Question banner showing the vote text. '
    'Stat boxes: Votes Cast, Expected, Turnout %, Majority (threshold). '
    'Control bar with status indicator + "▶ Open Voting", "■ Close Voting", "✅ Mark Complete" '
    'buttons. Live Results list below showing each answer with vote count and bar.')

body(doc, 'Managing the vote:')
data_table(doc,
    ['Button', 'Action', 'Notes'],
    [
        ['▶ Open Voting',    'Opens voting — members can now cast their vote',         'Disabled when voting is already open or no question is configured'],
        ['■ Close Voting',   'Closes voting',                                          'Disabled when voting is not open'],
        ['✅ Mark Complete', 'Finalises the vote permanently',                         'Disabled while voting is open; prompts for confirmation'],
        ['🔄 Reset & New Vote', 'Clears all votes and returns to an unconfigured state', 'Appears only after vote is marked complete'],
    ],
    col_widths=[3.8, 6.5, 5.7]
)

body(doc, (
    'When voting is open, voter phones show the congregational vote ballot automatically. '
    'When voting is closed or the vote is marked complete, voter phones update within 3 seconds.'
))

body(doc, 'The stat boxes update every 3 seconds:')
data_table(doc,
    ['Statistic', 'Meaning'],
    [
        ['Votes Cast',   'Total votes received (digital + paper)'],
        ['Expected',     'Expected voters (attending + absentee from Voting Setup)'],
        ['Turnout %',    'Votes Cast ÷ Expected × 100. Green when ≥ 50%.'],
        ['Majority',     'Minimum votes needed to pass, based on the configured threshold'],
    ],
    col_widths=[4.0, 12.0]
)

warning(doc, (
    'When the vote is marked complete, all voter phones show: '
    '"Thank you for your vote. The voting is now complete. '
    'The chairman will announce the results."'
))

h2(doc, '13b.  Voting Dashboard')
body(doc, (
    'The Voting Dashboard is a read-only view for the chairman showing live results '
    'for the congregational vote. It uses the same layout as the Election Dashboard. '
    'From the Voting Hub, click Voting Dashboard and enter the results password.'
))

screenshot(doc, 'Voting Dashboard',
    'Navy header "Voting Dashboard" with congregation name and meeting date. '
    'Status card showing vote status (Voting Open green / Voting Closed amber / Complete gold) '
    'and majority required. Info strip: Expected Voters and Absentee Votes pills. '
    'Vote question displayed in a bordered card. Stat boxes: Votes Cast, Expected, '
    'Turnout %, Majority. Answer rows below with vote counts, proportional bars, '
    'percentage of eligible voters, and green "✓ Majority" badge when threshold is reached.')

h2(doc, '13c.  Paper Ballot Entry — Congregational Vote')
body(doc, (
    'The same Paper Ballot Entry screen used for the election automatically switches to '
    'congregational vote mode when a congregational vote is active. '
    'Access it from the home screen using the paper ballot password.'
))

for step in [
    'The screen shows the vote question and the configured answer options.',
    'Click the member\'s selected answer to highlight it.',
    'If the ballot is from an absentee voter, tick the Absentee ballot checkbox.',
    'Click Submit Paper Vote. The vote is recorded and the form resets.',
]:
    numbered(doc, step)

screenshot(doc, 'Paper Ballot Entry — Congregational Vote',
    '"Congregational Vote" badge at top. Vote question in a bordered card. '
    'Answer options as selectable rows (one highlighted). Absentee ballot checkbox. '
    'Green "Record Paper Vote" button. Log of submitted paper votes below with '
    'Absentee badges on absentee entries.')


# ════════════════════════════════════════════════════════════════════════
# 14. BEST PRACTICES & SUGGESTED WORKFLOW
# ════════════════════════════════════════════════════════════════════════
h1(doc, '14.  Best Practices & Suggested Workflow')

h2(doc, 'Roles & Responsibilities')
data_table(doc,
    ['Role', 'Responsibilities'],
    [
        ['Election Officer',       'Operates Round Control and Voting Control; opens/closes voting; processes transitions; holds admin and voting passwords'],
        ['Chairman',               'Announces candidates, results, and voting instructions; accesses Election Dashboard and Voting Dashboard; holds results password'],
        ['Paper Ballot Volunteer', 'Operates Paper Ballot Entry; collects and enters paper ballots; holds paper ballot password'],
        ['Token Administrator',    'Generates and prints token cards; distributes them at the door; holds tokens password'],
        ['Door Steward',           'Distributes token cards to verified eligible voters as they enter'],
    ],
    col_widths=[4.0, 12.0]
)

h2(doc, 'Preparation — At Least One Week Before')
for item in [
    'Confirm nominees for each office with the consistory.',
    'Confirm the motion/resolution text to be voted on.',
    'Confirm the number of eligible voters.',
    'Test the full system end-to-end using a laptop and two phones.',
    'Change all passwords from their defaults. Record them securely.',
    'Print and cut token cards. Store securely.',
    'Print paper ballots for each office.',
]:
    bullet(doc, item)

h2(doc, 'Day of the Meeting — Before It Starts')
for step in [
    'Start the Python server: open Terminal, navigate to the election folder, run python3 server.py.',
    'Open http://localhost:8080/ in the admin browser. Enter the app password.',
    'Verify all election configuration: nominees, expected voters, absentee count, vote question.',
    'Confirm the voter URL and QR code are working — test on a phone.',
    'Set up the paper ballot station on a dedicated device with Paper Ballot Entry open.',
    'Distribute token cards at the entrance — one per eligible voter.',
]:
    numbered(doc, step)

h2(doc, 'Suggested Order of Events — Meeting Day')
for step in [
    'Chairman opens the meeting and introduces the election process.',
    'Chairman introduces the Elder nominees.',
    'Chairman asks members to open the voter URL or scan the QR code on their token card.',
    'Election officer opens Round Control → Start Round 1 Voting (Elder).',
    'Chairman: "Voting is now open. Please select your candidate(s) and submit."',
    'Paper ballot volunteer enters paper ballots during the voting period (absentee ballots first with checkbox ticked).',
    'Election officer monitors Ballots In and Participation %. When satisfied: Close Voting → End Round → confirm.',
    'Election officer processes Round Transition: confirm elected candidates (Step 1), set advancing candidates if needed (Step 2), votes per voter (Step 3).',
    'If positions filled: Complete This Office. If not: Launch Next Round (voting opens as closed — chairman announces round before officer opens voting).',
    'Chairman announces election results verbally (from Election Dashboard, not shown publicly).',
    'Repeat steps 2–10 for Deacon election if configured.',
    'After both offices complete: election officer sees Election Complete screen; voter phones show thank-you message.',
    'Chairman introduces the congregational vote motion.',
    'Election officer: Voting Hub → Voting Control → Open Voting.',
    'Chairman: "The motion is now open for voting. Please select your answer and cast your vote."',
    'Paper ballot volunteer enters paper votes during the voting period.',
    'Election officer monitors Votes Cast. When all votes are in: Close Voting → Mark Complete.',
    'Chairman accesses Voting Dashboard to review the result, then announces the outcome.',
    'Voter phones show: "Thank you for your vote. The voting is now complete. The chairman will announce the results."',
]:
    numbered(doc, step)

h2(doc, 'After the Meeting')
bullet(doc, 'Chairman may screenshot the Election Dashboard and Voting Dashboard for the consistory minutes.')
bullet(doc, 'When records are confirmed, use Reset All Election Data (Election Dashboard) to clear election data.')
bullet(doc, 'Use Reset & New Vote (Voting Control) to clear congregational vote data.')
bullet(doc, 'Stop the server: press Ctrl+C in the Terminal window.')
bullet(doc, 'Shred or store token cards securely.')

tip(doc, 'Always run a full test election before the actual meeting. Use at least two phones to verify the complete flow from token entry through to election complete.')
warning(doc, 'Never save Election Setup or Voting Setup during an active voting period. This resets all ballot data for that section.')


# ════════════════════════════════════════════════════════════════════════
# 15. COMPLETE MEETING WALKTHROUGH
# ════════════════════════════════════════════════════════════════════════
h1(doc, '15.  Complete Meeting Walkthrough')

body(doc, (
    'This section walks through a complete meeting: electing 2 Elders and 1 Deacon, '
    'followed by a congregational vote on a motion. '
    'Configuration: 80 attending voters, 5 absentee ballots. '
    'Round 1 majority = floor(85/2)+1 = 43. '
    'Round 2+ majority = floor(80/2)+1 = 41.'
))

h2(doc, 'Example Configuration')
data_table(doc,
    ['Setting', 'Elder', 'Deacon', 'Congregational Vote'],
    [
        ['Nominees / Question',      '4 nominees',   '3 nominees',   'Proposed resolution text'],
        ['Positions / Answer Options','2',            '1',            'In favour / Not in favour'],
        ['Votes per Voter',          '2 (Round 1)',   '1',            '—'],
        ['Expected Voters',          '80',           '80',           '80'],
        ['Absentee Votes',           '5',            '5',            '5'],
        ['Majority Threshold',       '43 (R1), 41 (R2+)', '43 (R1), 41 (R2+)', '50%+1 = 43'],
    ],
    col_widths=[4.5, 3.5, 3.0, 5.0]
)

h2(doc, 'Phase 1 — Setup (Day Before)')
for step in [
    'Run python3 server.py. Open admin hub, enter app password.',
    'Voter Tokens → enter tokens password → enter 85 → Generate Tokens.',
    'Election Setup → Settings → enter voter URL → Save URL. Change all passwords.',
    'Details tab: congregation name, date, Expected Voters = 80, Absentee = 5. Save.',
    'Elder tab: 4 nominees, Positions = 2, Votes = 2. Save. Print 15 Elder ballots.',
    'Deacon tab: 3 nominees, Positions = 1, Votes = 1. Save. Print 15 Deacon ballots.',
    'Voter Tokens → Print Token Cards. Cut into 85 cards.',
    'Voting Hub → Voting Setup: enter vote question, keep "In favour / Not in favour", Expected = 80, Absentee = 5, Threshold = 50%+1. Save.',
    'Test voter URL on a phone. Confirm token screen appears.',
]:
    numbered(doc, step)

h2(doc, 'Phase 2 — Elder Election, Round 1')
for step in [
    'Chairman introduces Elder nominees.',
    'Round Control → Start Round 1 Voting.',
    'Members enter tokens and select up to 2 candidates.',
    'Paper ballot volunteer enters 5 absentee ballots (Absentee checkbox ticked), then remaining paper ballots.',
    'Election officer: Close Voting → End Round (confirm) → process Round Transition.',
    'Scenario: 2 candidates reach majority. Complete This Office. Chairman announces elected Elders.',
]:
    numbered(doc, step)

h2(doc, 'Phase 3 — Deacon Election (may require 2 rounds)')
for step in [
    'Round Control now shows "Ready to Start" for Deacon. Chairman introduces nominees.',
    'Start Round 1 Voting. Members re-enter tokens for Deacon.',
    'If no majority in Round 1: Launch Next Round (exclude lowest candidate in Step 2).',
    'Round 2: majority reached. Complete This Office. Chairman announces elected Deacon.',
    'Voter phones now show thank-you message. Election officer sees Election Complete screen.',
]:
    numbered(doc, step)

h2(doc, 'Phase 4 — Congregational Vote')
for step in [
    'Chairman introduces the motion and reads it to the congregation.',
    'Voting Hub → Voting Control → Open Voting.',
    'Voter phones automatically switch to the congregational vote ballot (token already set — no re-entry needed for those who already voted in the election).',
    'Chairman: "Please vote on the motion using your token card."',
    'Paper ballot volunteer enters paper votes (switches to vote mode automatically).',
    'Election officer: Close Voting → Mark Complete (confirm).',
    'Chairman accesses Voting Dashboard with results password. Announces outcome.',
    'Voter phones show: "Thank you for your vote. The voting is now complete. The chairman will announce the results."',
]:
    numbered(doc, step)


# ════════════════════════════════════════════════════════════════════════
# 16. TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════
h1(doc, '16.  Troubleshooting')

data_table(doc,
    ['Problem', 'Likely Cause', 'Solution'],
    [
        ['App password screen does not dismiss',
         'Wrong password entered',
         'Re-enter the correct app password. Default is votevote2024 (change it first!).'],
        ['Voter page shows "No Voting Configured"',
         'server.py not running, or setup not saved while server was running',
         'Confirm server.py is running in Terminal. Re-open Election Setup and save again.'],
        ['Voter devices cannot reach vote.html',
         'Devices not on same Wi-Fi, wrong URL, or server not running',
         'Confirm all devices are on the same network. Check voter URL in Settings. Restart server.py if needed.'],
        ['"Failed to record vote" on phone',
         'Server stopped or phone lost network',
         'Restart server.py in Terminal. Ensure phone is on correct Wi-Fi.'],
        ['"Invalid token" error',
         'Code mistyped or tokens were regenerated after printing',
         'Double-check the 4-digit code. If tokens were regenerated, all old cards are invalid.'],
        ['"Token already used" error',
         'Voter already voted this round, or paper ballot submitted for their token',
         'The voter has already voted this round. Check with the chairman if in error.'],
        ['Voter phone stuck on "Voting Closed"',
         'Voting closed in Round Control; voter was already on ballot',
         'Normal behaviour — tell voter to wait. Page updates automatically when next round opens.'],
        ['Vote counts not updating in Round Control',
         'Browser tab lost network connection',
         'Refresh the browser. Counts update every 3 seconds when server is reachable.'],
        ['Round Control shows wrong office after completion',
         'Other office not configured (no nominees saved)',
         'Go to Election Setup → configure the other office, or proceed if it was intentionally skipped.'],
        ['Paper Ballot Entry form not appearing',
         'Voting is closed or no voting is active',
         'Ensure voting is open in Round Control or Voting Control. Screen updates within 4 seconds.'],
        ['Congregational vote ballot not showing on voter phone',
         'Voting Control: voting is not open',
         'Go to Voting Control → Open Voting. Voter phones update within 3 seconds.'],
        ['After election completes, vote ballot not appearing',
         'Congregational vote not yet open in Voting Control',
         'Open Voting Control and click Open Voting. Phones transition automatically.'],
        ['Login fails on mobile (password rejected)',
         'Mobile keyboard autocapitalising or adding space',
         'Type the password carefully. Password fields disable autocorrect. Avoid copy-pasting.'],
        ['Forgotten admin or voting password',
         'Password changed and lost',
         'No recovery option. Use Reset All Election Data — this deletes everything. Always record passwords.'],
        ['Paper ballot station shows wrong election type',
         'Both election and congregational vote are active simultaneously',
         'Only one type of voting can be active at a time. Close the election or vote before opening the other.'],
    ],
    col_widths=[4.2, 4.3, 7.5]
)


# ════════════════════════════════════════════════════════════════════════
# 17. QUICK REFERENCE
# ════════════════════════════════════════════════════════════════════════
h1(doc, '17.  Quick Reference')

h2(doc, 'All Passwords (fill in your custom passwords before the meeting)')
data_table(doc,
    ['Password', 'Default', 'Protects', 'Custom Password'],
    [
        ['App (Landing)',    'votevote2024',   'Entire admin hub — session login',          '____________________'],
        ['Admin',           'election2024',   'Election Setup, Round Control',             '____________________'],
        ['Voting',          'voting2024',     'Voting Setup, Voting Control',              '____________________'],
        ['Results',         'results2024',    'Election Dashboard, Voting Dashboard',      '____________________'],
        ['Tokens',          'tokens2024',     'Voter Tokens section',                      '____________________'],
        ['Paper Ballot',    'paperentry2024', 'Paper Ballot Entry station',                '____________________'],
    ],
    col_widths=[2.8, 3.0, 5.2, 5.0]
)

h2(doc, 'Voter URL & Network')
data_table(doc,
    ['', 'Value'],
    [
        ['Voter Page URL',    '________________________________________________'],
        ['Server IP Address', '________________________________________________'],
    ],
    col_widths=[4.5, 11.5]
)

h2(doc, 'Majority Thresholds')
data_table(doc,
    ['Setting', 'Value'],
    [
        ['Expected Voters (attending)',           '________'],
        ['Absentee Votes',                        '________'],
        ['Round 1 Total (attending + absentee)',  '________'],
        ['Round 1 Majority Required',             'floor(R1 Total / 2) + 1  =  ________'],
        ['Round 2+ Majority Required',            'floor(Attending / 2) + 1  =  ________'],
        ['Congregational Vote Majority',          '____________________________'],
    ],
    col_widths=[6.0, 10.0]
)

h2(doc, 'Screen Navigation Summary')
data_table(doc,
    ['Screen', 'Who', 'Password', 'Purpose'],
    [
        ['Home Screen',              'All officers',     'App password (session)',  'Navigation hub, session entry point'],
        ['Election Setup',           'Election Officer', 'Admin',                   'Configure nominees, tokens, passwords, URL'],
        ['Round Control',            'Election Officer', 'Admin (hub)',             'Open/close voting, monitor ballots, transitions'],
        ['Voter Tokens',             'Token Admin',      'Tokens',                  'Generate tokens, print cards, QR code'],
        ['Voting Setup',             'Election Officer', 'Voting (hub)',            'Configure vote question, answers, majority'],
        ['Voting Control',           'Election Officer', 'Voting (hub)',            'Open/close congregational voting, monitor results'],
        ['Paper Ballot Entry',       'Volunteer',        'Paper Ballot',            'Enter paper votes — auto-detects election or vote mode'],
        ['Election Dashboard',       'Chairman',         'Results',                 'Live election status and vote counts (read-only)'],
        ['Voting Dashboard',         'Chairman',         'Results',                 'Live congregational vote results (read-only)'],
        ['Election Complete',        'Election Officer', 'None (auto)',             'Shows elected candidates after all offices done'],
        ['Voter Page (vote.html)',   'All voters',       'Token code',             'Submit digital election ballot and congregational vote'],
    ],
    col_widths=[3.8, 3.0, 3.2, 6.0]
)

h2(doc, 'Election Day Checklist')
checklist = [
    'server.py running in Terminal',
    'App password changed from default and recorded',
    'Admin password changed from default and recorded',
    'Voting password changed from default and recorded',
    'Results password changed from default and shared with chairman only',
    'Tokens password changed from default and recorded',
    'Paper ballot password changed from default and shared with volunteer',
    'Voter URL saved in Settings and tested on a phone',
    'QR code scanned and confirmed to open vote.html',
    'Congregation name, meeting date, Expected Voters, Absentee Votes saved in Details tab',
    'Elder nominees, positions, votes per voter saved (if Elder election)',
    'Deacon nominees, positions, votes per voter saved (if Deacon election)',
    'Congregational vote question, answers, expected voters, majority threshold saved',
    'Voter tokens generated — token cards printed, cut, and ready to distribute',
    'Paper ballots printed for each office',
    'Token cards distributed to eligible voters at the entrance',
    'Paper ballot station set up on a dedicated device with volunteer ready',
    'Paper Ballot Entry open on volunteer\'s device (logged in)',
    'Chairman briefed on results password and how to access both dashboards',
]
for item in checklist:
    p = doc.add_paragraph()
    para_space(p, before=0, after=4)
    r = p.add_run('☐   ' + item)
    r.font.size = Pt(11)

# Footer
footer_p = doc.add_paragraph()
para_space(footer_p, before=24, after=0)
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(
    'Church Voting App — Complete User Manual  |  Version 2.6.1\n'
    'All data is stored on the local server. No votes leave the meeting network.'
)
fr.font.size = Pt(9)
fr.font.color.rgb = GREY

# ── Save ──────────────────────────────────────────────────────────────────────
import os
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'manual.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
